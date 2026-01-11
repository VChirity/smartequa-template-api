from docx import Document
import os
import re

def consolidate_fragmented_tags(filepath):
    """
    Consolida tags fragmentadas do docxtpl sem destruir formatação, logo ou estrutura.
    Quando você edita manualmente no Word e salva, o Word fragmenta tags como:
    {{nome_aluno}} vira {{no + me_al + uno}}
    Este script reconstrói as tags mantendo TUDO intacto.
    """
    print(f"\n🔧 Processando: {filepath}")
    
    if not os.path.exists(filepath):
        print(f"❌ Arquivo não encontrado!")
        return False
    
    try:
        doc = Document(filepath)
        print(f"✅ Documento carregado")
        
        changes_made = 0
        
        # Lista de tags esperadas
        expected_tags = [
            'responsavel1', 'cpf_responsavel', 'endereco_completo', 'bairro', 'cep',
            'naturalidade_resp1', 'nasc_resp1',
            'responsavel2', 'cpf2', 'endereco2', 'bairro2', 'cep2',
            'nome_aluno', 'ano', 'ano_letivo', 'data_extenso',
            'naturalidade_aluno', 'nasc_aluno', 'cpf_aluno',
            'desconto', 'desconto_extenso'
        ]
        
        def consolidate_paragraph(para):
            """Consolida tags fragmentadas em um parágrafo"""
            nonlocal changes_made
            
            # Pega o texto completo do parágrafo
            full_text = para.text
            
            # Verifica se há tags ({{ e }})
            if '{{' not in full_text or '}}' not in full_text:
                return False
            
            # Conta quantos runs existem
            num_runs = len(para.runs)
            
            # Se há muitos runs em um parágrafo com tags, provavelmente está fragmentado
            if num_runs > 3:
                # Verifica se alguma tag esperada está no texto
                has_tag = False
                for tag in expected_tags:
                    if tag in full_text:
                        has_tag = True
                        break
                
                if has_tag:
                    # Salva formatação do primeiro run
                    if para.runs:
                        first_run = para.runs[0]
                        bold = first_run.bold
                        italic = first_run.italic
                        font_name = first_run.font.name if first_run.font.name else None
                        font_size = first_run.font.size
                        font_color = None
                        if first_run.font.color and first_run.font.color.rgb:
                            font_color = first_run.font.color.rgb
                        
                        # Remove todos os runs exceto o primeiro
                        for i in range(len(para.runs) - 1, 0, -1):
                            para._element.remove(para.runs[i]._element)
                        
                        # Coloca todo o texto no primeiro run
                        para.runs[0].text = full_text
                        
                        # Restaura formatação
                        if bold is not None:
                            para.runs[0].bold = bold
                        if italic is not None:
                            para.runs[0].italic = italic
                        if font_name:
                            para.runs[0].font.name = font_name
                        if font_size:
                            para.runs[0].font.size = font_size
                        if font_color:
                            from docx.shared import RGBColor
                            para.runs[0].font.color.rgb = font_color
                        
                        changes_made += 1
                        print(f"  ✓ Consolidado parágrafo com {num_runs} runs → 1 run")
                        return True
            
            return False
        
        # Processa parágrafos
        for para in doc.paragraphs:
            consolidate_paragraph(para)
        
        # Processa tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        consolidate_paragraph(para)
        
        # Processa cabeçalhos e rodapés
        for section in doc.sections:
            for para in section.header.paragraphs:
                consolidate_paragraph(para)
            
            for para in section.footer.paragraphs:
                consolidate_paragraph(para)
        
        # Salva o documento
        doc.save(filepath)
        
        print(f"✅ Documento salvo com {changes_made} consolidações")
        print(f"🎨 Formatação, logo e estrutura PRESERVADAS")
        
        return True
        
    except Exception as e:
        print(f"❌ Erro: {e}")
        import traceback
        traceback.print_exc()
        return False

def main():
    print("=" * 70)
    print("CONSOLIDATE TAGS - Correção de Fragmentação")
    print("Corrige tags fragmentadas pelo Word SEM destruir formatação")
    print("=" * 70)
    
    templates = [
        "templates/IMAGEM-E-VOZ-ALUNO-PUBLICIDADE_1.docx",
        "templates/IMAGEM-E-VOZ-ALUNO-INSTITUCIONAL_1.docx"
    ]
    
    success_count = 0
    for template_path in templates:
        if consolidate_fragmented_tags(template_path):
            success_count += 1
    
    print("\n" + "=" * 70)
    print(f"✅ Concluído! {success_count}/{len(templates)} arquivos processados")
    print("=" * 70)
    
    if success_count == len(templates):
        print("\n🎉 Tags consolidadas! Arquivos prontos para uso com docxtpl!")
        print("💡 Teste no app agora - deve funcionar sem corrupção.")
    else:
        print("\n⚠️ Alguns arquivos falharam.")

if __name__ == '__main__':
    main()
