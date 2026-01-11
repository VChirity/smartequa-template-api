from docx import Document
import re

def create_discount_version():
    source = 'templates/template_contrato2025_2.docx'
    dest = 'templates/template_contratoDESCONTO2025_2.docx'
    
    print(f"Lendo base: {source}...")
    try:
        doc = Document(source)
    except:
        # Fallback para assets se não achar em templates
        doc = Document('assets/template_contrato2025_2.docx')

    # Texto da Cláusula de Desconto (Tags já corrigidas)
    new_clause_text = (
        "CLÁUSULA 3ª - A contratada concede por mera liberalidade e para o ano letivo vigente, "
        "exclusivamente, um desconto no valor da anuidade devida no percentual de {{desconto}}% "
        "({{desconto_extenso}}), desde que a mesma seja quitada na data de seu respectivo vencimento, "
        "não implicando tal desconto em novação da anualidade devida.\n"
        "PARÁGRAFO ÚNICO - No caso de atraso de pagamento, o desconto será suprimido do saldo "
        "remanescente das mensalidades até que haja a regularização das parcelas em aberto, "
        "sem o prejuízo das demais cláusulas contratuais. Após a regularização das parcelas em aberto, "
        "o desconto poderá retornar incidindo única e exclusivamente sobre o saldo remanescente "
        "da anuidade a partir da parcela do mês subsequente a aquela data."
    )

    # 1. Renumeração (De trás para frente para evitar conflitos, ex: 14->15, ..., 3->4)
    # Vamos varrer e substituir números de cláusulas
    # Regex procura por "CLÁUSULA Xª"
    for p in doc.paragraphs:
        if p.text.strip().upper().startswith("CLÁUSULA"):
            # Extrai o número
            match = re.search(r"CLÁUSULA\s+(\d+)ª", p.text, re.IGNORECASE)
            if match:
                num = int(match.group(1))
                if num >= 3:
                    # Incrementa +1 (3 vira 4)
                    p.text = p.text.replace(f"{num}ª", f"{num+1}ª")

    # 2. Inserção da Nova Cláusula 3ª
    # Procura onde ficou a "CLÁUSULA 4ª" (antiga 3ª) para inserir antes dela
    for i, p in enumerate(doc.paragraphs):
        if "CLÁUSULA 4ª" in p.text.upper():
            # Insere o novo parágrafo antes deste
            new_p = p.insert_paragraph_before(new_clause_text)
            # Tenta manter estilo (opcional, assume estilo padrão do doc)
            break

    doc.save(dest)
    print(f"✅ Sucesso! Arquivo {dest} recriado com cláusula de desconto e renumeração.")

if __name__ == '__main__':
    create_discount_version()
