from docx import Document

# Carregar o template atual
template_path = 'templates_quadros/notas/quadro_notas_template.docx'

print(f"Carregando template: {template_path}")
doc = Document(template_path)

# Procurar em todas as tabelas
print(f"\nTotal de tabelas: {len(doc.tables)}")

encontrado = False
for t_idx, table in enumerate(doc.tables):
    print(f"\nTabela {t_idx}: {len(table.rows)} linhas x {len(table.columns)} colunas")
    
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            texto_cell = cell.text
            
            # Procurar pela célula com checkboxes
            if '☐' in texto_cell or 'Nota única' in texto_cell:
                print(f"\n✓ Encontrado na Tabela {t_idx}, Linha {r_idx}, Célula {c_idx}")
                print(f"Texto atual: {texto_cell[:200]}")
                
                # Procurar o parágrafo específico com os checkboxes
                for p_idx, para in enumerate(cell.paragraphs):
                    if '☐' in para.text or 'Nota única' in para.text:
                        print(f"  Parágrafo {p_idx}: {para.text[:100]}")
                        
                        # Limpar e adicionar novo texto com lógica condicional
                        para.clear()
                        novo_texto = (
                            "{% if tipo_calculo == 'Nota única' %}☑{% else %}☐{% endif %}Nota única"
                            "    "
                            "{% if tipo_calculo == 'Média Aritmética' %}☑{% else %}☐{% endif %}Média Aritmética"
                            "    "
                            "{% if tipo_calculo == 'Adição das notas' %}☑{% else %}☐{% endif %}Adição das notas"
                            "    "
                            "{% if tipo_calculo == 'Outro' %}☑{% else %}☐{% endif %}Outro (ESCREVER ABAIXO)"
                        )
                        
                        para.add_run(novo_texto)
                        print(f"  ✓ Atualizado!")
                        encontrado = True
                        break
                
                if encontrado:
                    break
        if encontrado:
            break
    if encontrado:
        break

if encontrado:
    # Salvar
    output_path = 'templates_quadros/notas/quadro_notas_template.docx'
    doc.save(output_path)
    print(f"\n✅ Template atualizado com checkboxes condicionais!")
else:
    print("\n❌ Não encontrei os checkboxes!")
