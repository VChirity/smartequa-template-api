from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Carregar template
template_path = 'templates_quadros/notas/quadro_notas_template.docx'
print(f"Carregando template: {template_path}")
doc = Document(template_path)

# 1. Centralizar número do loop na tabela de notas
print("\n1. Centralizando número do loop...")
tabela_notas = doc.tables[0]

# Encontrar linha de dados
for i, row in enumerate(tabela_notas.rows):
    texto = ' '.join([cell.text for cell in row.cells])
    if '{{loop.index}}' in texto:
        print(f"   Linha de dados encontrada: {i}")
        # Centralizar primeira célula (número)
        cell = row.cells[0]
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        print("   ✓ Número centralizado")
        break

# 2. Evitar quebra de página entre tabelas
print("\n2. Configurando para evitar quebra de página feia...")

# Adicionar propriedade "keep together" nas tabelas
for idx, table in enumerate(doc.tables):
    print(f"   Tabela {idx}:")
    
    # Adicionar propriedade para manter linhas juntas
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    # Adicionar keep together
    keepTogether = OxmlElement('w:cantSplit')
    tblPr.append(keepTogether)
    print("     ✓ Keep together adicionado")

# 3. Adicionar propriedade "keep with next" nos parágrafos antes das tabelas
print("\n3. Mantendo parágrafos com tabelas seguintes...")
for i, para in enumerate(doc.paragraphs):
    # Adicionar keep with next para evitar quebra
    pPr = para._element.get_or_add_pPr()
    keepNext = OxmlElement('w:keepNext')
    pPr.append(keepNext)
    print(f"   ✓ Parágrafo {i} configurado")

# Salvar
output_path = 'templates_quadros/notas/quadro_notas_template.docx'
doc.save(output_path)

print(f"\n✅ Template corrigido!")
print("   - Número do loop centralizado")
print("   - Tabelas configuradas para não quebrar")
print("   - Parágrafos mantidos com próxima tabela")
