from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_image_term_from_contract(base_template, dest_path, title, content_paragraphs):
    """
    Cria um termo de imagem a partir do template de contrato,
    preservando o parágrafo 0 (logo/cabeçalho) e as seções.
    """
    print(f"\n🔧 Criando: {os.path.basename(dest_path)}")
    
    if not os.path.exists(base_template):
        print(f"❌ Template base não encontrado: {base_template}")
        return False
    
    try:
        # Carrega o template base (que tem logo e formatação corretas)
        print("📂 Carregando template base com logo...")
        doc = Document(base_template)
        
        # Preserva o parágrafo 0 (onde a logo está ancorada)
        print("🖼️ Preservando parágrafo 0 (logo/cabeçalho)...")
        
        # Remove todos os parágrafos EXCETO o primeiro (índice 0)
        print("🗑️ Removendo conteúdo antigo (preservando cabeçalho)...")
        paragraphs_to_remove = list(doc.paragraphs[1:])
        for para in paragraphs_to_remove:
            p = para._element
            p.getparent().remove(p)
        
        # Adiciona o título
        print("📝 Adicionando título...")
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(14)
        title_run.font.name = 'Calibri'
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Adiciona linha em branco
        doc.add_paragraph()
        
        # Adiciona os parágrafos do conteúdo
        print("✏️ Adicionando conteúdo...")
        for content in content_paragraphs:
            para = doc.add_paragraph()
            run = para.add_run(content)
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0, 0, 0)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Salva o documento
        print(f"💾 Salvando: {dest_path}")
        doc.save(dest_path)
        
        print(f"✅ Termo criado com sucesso!")
        print(f"🎨 Logo e estrutura preservadas")
        
        return True
        
    except Exception as e:
        print(f"❌ Erro: {e}")
        import traceback
        traceback.print_exc()
        return False

def main():
    print("=" * 70)
    print("SMART CREATE TERMS - Criação Inteligente de Termos")
    print("Preserva logo e estrutura do contrato base")
    print("=" * 70)
    
    base_template = "templates/template_contrato2025_2.docx"
    
    # Conteúdo do Termo de Publicidade
    publicidade_content = [
        "Eu, {{responsavel1}}, {{naturalidade_resp1}}, nascido(a) em {{nasc_resp1}}, inscrito(a) no CPF/MF sob nº {{cpf_responsavel}}, residente no endereço {{endereco_completo}}, Rio de Janeiro – RJ, responsável pelo/pela criança/adolescente {{nome_aluno}}, {{naturalidade_aluno}}, nascido(a) em {{nasc_aluno}}, inscrito(a) no CPF/MF sob nº {{cpf_aluno}}, pelo presente instrumento, AUTORIZO o CURSO DE ESPECIALIZAÇÃO EQUAÇÃO LTDA, com sede na Rua Mendes Tavares, nº108, Vila Isabel, Rio de Janeiro – RJ, inscrita no CNPJ/MF sob o nº 42.319.202.001-40, a fazer uso da imagem e/ou voz do menor acima identificado, em todo e qualquer material entre fotos, documentos e outros meios de comunicação, para campanhas publicitárias, sejam essas destinadas à divulgação ao público em geral e/ou apenas para uso desta escola.",
        "",
        "A presente autorização é concedida a título gratuito, abrangendo o uso da imagem acima mencionada em todo território nacional e no exterior, sob qualquer forma e meios, ou sejam, em destaque: (I) outdoor; (II) busdoor; folhetos em geral (encartes, mala direta, catálogo, etc.); (III) folder de apresentação; (IV) anúncios em revistas e jornais em geral; (V) home page; (VI) cartazes; (VII) backlight; (VIII) mídia eletrônica (internet, painéis, vídeotapes, televisão, cinema, programa para rádio, entre outros).",
        "",
        "Por esta ser a expressão da minha vontade, declaro que autorizo o uso acima descrito sem que nada haja a ser reclamado, e assino a presente autorização em 02 (duas) vias de igual teor e forma.",
        "",
        "Rio de Janeiro, {{data_extenso}}.",
        "",
        "",
        "________________________________________________",
        "{{responsavel1}}",
        "CPF: {{cpf_responsavel}}"
    ]
    
    # Conteúdo do Termo Institucional (mesma coisa mas "campanhas institucionais")
    institucional_content = [
        "Eu, {{responsavel1}}, {{naturalidade_resp1}}, nascido(a) em {{nasc_resp1}}, inscrito(a) no CPF/MF sob nº {{cpf_responsavel}}, residente no endereço {{endereco_completo}}, Rio de Janeiro – RJ, responsável pelo/pela criança/adolescente {{nome_aluno}}, {{naturalidade_aluno}}, nascido(a) em {{nasc_aluno}}, inscrito(a) no CPF/MF sob nº {{cpf_aluno}}, pelo presente instrumento, AUTORIZO o CURSO DE ESPECIALIZAÇÃO EQUAÇÃO LTDA, com sede na Rua Mendes Tavares, nº108, Vila Isabel, Rio de Janeiro – RJ, inscrita no CNPJ/MF sob o nº 42.319.202.001-40, a fazer uso da imagem e/ou voz do menor acima identificado, em todo e qualquer material entre fotos, documentos e outros meios de comunicação, para campanhas institucionais, sejam essas destinadas à divulgação ao público em geral e/ou apenas para uso desta escola.",
        "",
        "A presente autorização é concedida a título gratuito, abrangendo o uso da imagem acima mencionada em todo território nacional e no exterior, sob qualquer forma e meios, ou sejam, em destaque: (I) outdoor; (II) busdoor; folhetos em geral (encartes, mala direta, catálogo, etc.); (III) folder de apresentação; (IV) anúncios em revistas e jornais em geral; (V) home page; (VI) cartazes; (VII) backlight; (VIII) mídia eletrônica (internet, painéis, vídeotapes, televisão, cinema, programa para rádio, entre outros).",
        "",
        "Por esta ser a expressão da minha vontade, declaro que autorizo o uso acima descrito sem que nada haja a ser reclamado, e assino a presente autorização em 02 (duas) vias de igual teor e forma.",
        "",
        "Rio de Janeiro, {{data_extenso}}.",
        "",
        "",
        "________________________________________________",
        "{{responsavel1}}",
        "CPF: {{cpf_responsavel}}"
    ]
    
    templates = [
        {
            'dest': 'templates/IMAGEM-E-VOZ-ALUNO-PUBLICIDADE_1.docx',
            'title': 'TERMO DE AUTORIZAÇÃO DE USO\nIMAGEM E VOZ DE ALUNO - PUBLICIDADE',
            'content': publicidade_content
        },
        {
            'dest': 'templates/IMAGEM-E-VOZ-ALUNO-INSTITUCIONAL_1.docx',
            'title': 'TERMO DE AUTORIZAÇÃO DE USO\nIMAGEM E VOZ DE ALUNO - INSTITUCIONAL',
            'content': institucional_content
        }
    ]
    
    success_count = 0
    for template in templates:
        if create_image_term_from_contract(base_template, template['dest'], template['title'], template['content']):
            success_count += 1
    
    print("\n" + "=" * 70)
    print(f"✅ Concluído! {success_count}/{len(templates)} termos criados")
    print("=" * 70)
    
    if success_count == len(templates):
        print("\n🎉 Termos de Imagem criados com logo e estrutura preservadas!")
        print("💡 Baseados no template de contrato que funciona")
        print("🧪 Teste no app agora - deve funcionar sem corrupção!")
    else:
        print("\n⚠️ Alguns termos falharam.")

if __name__ == '__main__':
    main()
