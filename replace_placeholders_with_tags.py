from docx import Document
import os

def replace_placeholders_with_tags(filepath):
    """
    Substitui os placeholders em vermelho pelas tags do docxtpl em preto
    """
    print(f"\n🔧 Processando: {filepath}")
    
    if not os.path.exists(filepath):
        print(f"❌ Arquivo não encontrado!")
        return False
    
    try:
        doc = Document(filepath)
        print(f"✅ Documento carregado")
        
        # Mapeamento de placeholders para tags
        replacements = {
            '(NOME COMPLETO DO RESPONSÁVEL)': '{{responsavel1}}',
            '(NATURALIDADE_RESP1)': '{{naturalidade_resp1}}',
            '(DATA DE NASCIMENTO DO RESPONSÁVEL)': '{{nasc_resp1}}',
            '(CPF_RESPONSAVEL)': '{{cpf_responsavel}}',
            '{{cpf_responsavel}}': '{{cpf_responsavel}}',  # Já está correto
            '(ENDEREÇO COMPLETO DO RESPONSÁVEL)': '{{endereco_completo}}',
            '(NOME COMPLETO DO ALUNO)': '{{nome_aluno}}',
            '(NATURALIDADE DO ALUNO)': '{{naturalidade_aluno}}',
            '(DATA DE NASCIMENTO DO ALUNO)': '{{nasc_aluno}}',
            '(CPF DO ALUNO)': '{{cpf_aluno}}',
            '{{cpf_aluno}}': '{{cpf_aluno}}',  # Já está correto
        }
        
        changes_made = 0
        
        # Processa todos os parágrafos
        for para in doc.paragraphs:
            for run in para.runs:
                original_text = run.text
                new_text = original_text
                
                # Substitui cada placeholder
                for placeholder, tag in replacements.items():
                    if placeholder in new_text:
                        new_text = new_text.replace(placeholder, tag)
                        changes_made += 1
                        print(f"  ✓ Substituído: {placeholder} → {tag}")
                
                # Se houve mudança, atualiza o texto e garante cor preta
                if new_text != original_text:
                    run.text = new_text
                    # Define cor como preta (RGB 0,0,0)
                    if run.font.color:
                        run.font.color.rgb = None  # Remove cor customizada (volta ao preto padrão)
        
        # Processa tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            original_text = run.text
                            new_text = original_text
                            
                            for placeholder, tag in replacements.items():
                                if placeholder in new_text:
                                    new_text = new_text.replace(placeholder, tag)
                                    changes_made += 1
                                    print(f"  ✓ Substituído: {placeholder} → {tag}")
                            
                            if new_text != original_text:
                                run.text = new_text
                                if run.font.color:
                                    run.font.color.rgb = None
        
        # Processa cabeçalhos e rodapés
        for section in doc.sections:
            for para in section.header.paragraphs:
                for run in para.runs:
                    original_text = run.text
                    new_text = original_text
                    
                    for placeholder, tag in replacements.items():
                        if placeholder in new_text:
                            new_text = new_text.replace(placeholder, tag)
                            changes_made += 1
                            print(f"  ✓ Substituído: {placeholder} → {tag}")
                    
                    if new_text != original_text:
                        run.text = new_text
                        if run.font.color:
                            run.font.color.rgb = None
            
            for para in section.footer.paragraphs:
                for run in para.runs:
                    original_text = run.text
                    new_text = original_text
                    
                    for placeholder, tag in replacements.items():
                        if placeholder in new_text:
                            new_text = new_text.replace(placeholder, tag)
                            changes_made += 1
                            print(f"  ✓ Substituído: {placeholder} → {tag}")
                    
                    if new_text != original_text:
                        run.text = new_text
                        if run.font.color:
                            run.font.color.rgb = None
        
        # Salva o documento
        doc.save(filepath)
        
        print(f"✅ Documento salvo com {changes_made} substituições")
        print(f"🎨 Tags em PRETO, formatação preservada")
        
        return True
        
    except Exception as e:
        print(f"❌ Erro: {e}")
        import traceback
        traceback.print_exc()
        return False

def main():
    print("=" * 70)
    print("REPLACE PLACEHOLDERS WITH TAGS")
    print("Substituindo placeholders por tags docxtpl em PRETO")
    print("=" * 70)
    
    templates = [
        "templates/IMAGEM-E-VOZ-ALUNO-PUBLICIDADE_1.docx",
        "templates/IMAGEM-E-VOZ-ALUNO-INSTITUCIONAL_1.docx"
    ]
    
    success_count = 0
    for template_path in templates:
        if replace_placeholders_with_tags(template_path):
            success_count += 1
    
    print("\n" + "=" * 70)
    print(f"✅ Concluído! {success_count}/{len(templates)} arquivos processados")
    print("=" * 70)
    
    if success_count == len(templates):
        print("\n🎉 Tags inseridas em PRETO nos 2 arquivos!")
        print("💡 Abra os arquivos no Word para verificar.")
    else:
        print("\n⚠️ Alguns arquivos falharam.")

if __name__ == '__main__':
    main()
