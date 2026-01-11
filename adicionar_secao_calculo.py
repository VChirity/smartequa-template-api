from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_background(cell, color):
    """Definir cor de fundo da célula"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

# Carregar template
template_path = 'templates_quadros/notas/quadro_notas_template.docx'
print(f"Carregando template: {template_path}")
doc = Document(template_path)

# Adicionar espaço após a tabela principal
doc.add_paragraph()

# Adicionar tabela de cálculo da média
print("Adicionando seção CÁLCULO DA MÉDIA...")
table_calculo = doc.add_table(rows=3, cols=1)

# Linha 0: Título "CÁLCULO DA MÉDIA"
row0 = table_calculo.rows[0]
cell0 = row0.cells[0]
cell0.text = 'CÁLCULO DA MÉDIA'
cell0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
run = cell0.paragraphs[0].runs[0]
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(255, 255, 255)
set_cell_background(cell0, '1F4E78')  # Azul escuro

# Linha 1: Checkboxes com lógica condicional
row1 = table_calculo.rows[1]
cell1 = row1.cells[0]
set_cell_background(cell1, '9DC3E6')  # Azul claro

# Adicionar texto com checkboxes condicionais
para1 = cell1.paragraphs[0]
texto_checkboxes = (
    "{% if tipo_calculo == 'Nota única' %}☑{% else %}☐{% endif %}Nota única"
    "            "
    "{% if tipo_calculo == 'Média Aritmética' %}☑{% else %}☐{% endif %}Média Aritmética"
    "            "
    "{% if tipo_calculo == 'Adição das notas' %}☑{% else %}☐{% endif %}Adição das notas"
    "            "
    "{% if tipo_calculo == 'Outro' %}☑{% else %}☐{% endif %}Outro (ESCREVER ABAIXO)"
)
para1.text = texto_checkboxes
para1.alignment = WD_ALIGN_PARAGRAPH.LEFT

# Linha 2: Explicação do cálculo
row2 = table_calculo.rows[2]
cell2 = row2.cells[0]
set_cell_background(cell2, 'D9E1F2')  # Azul muito claro
para2 = cell2.paragraphs[0]
para2.text = 'Como calcular a média: {{explicacao_calculo}}'
para2.alignment = WD_ALIGN_PARAGRAPH.LEFT

print("✓ Seção CÁLCULO DA MÉDIA adicionada")

# Adicionar espaço
doc.add_paragraph()

# Adicionar rodapé
print("Adicionando rodapé...")
para_prof = doc.add_paragraph()
para_prof.text = 'Professor: {{professor}}'

para_disc = doc.add_paragraph()
para_disc.text = 'Disciplina: {{disciplina}}'

# Aplicar cor amarela ao rodapé
for para in [para_prof, para_disc]:
    for run in para.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    # Adicionar fundo amarelo ao parágrafo
    pPr = para._element.get_or_add_pPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'FFC000')
    pPr.append(shading)

print("✓ Rodapé adicionado")

# Salvar
output_path = 'templates_quadros/notas/quadro_notas_template.docx'
doc.save(output_path)

print(f"\n✅ Template atualizado com seção de cálculo!")
print("Teste novamente no SmartEqua!")
