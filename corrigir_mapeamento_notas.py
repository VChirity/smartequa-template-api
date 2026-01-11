from docx import Document

# Carregar template
template_path = 'templates_quadros/notas/quadro_notas_template.docx'
print(f"Carregando template: {template_path}")
doc = Document(template_path)

# Encontrar a tabela de notas
tabela_notas = None
for table in doc.tables:
    if len(table.rows) >= 5:
        tabela_notas = table
        break

if not tabela_notas:
    print("ERRO: Tabela não encontrada!")
    exit(1)

print(f"Tabela encontrada com {len(tabela_notas.rows)} linhas")

# Encontrar a linha de dados (linha 5, índice 5)
linha_dados = None
for i, row in enumerate(tabela_notas.rows):
    texto_linha = ' '.join([cell.text for cell in row.cells])
    if '{{aluno.nome}}' in texto_linha:
        linha_dados = i
        print(f"Linha de dados encontrada: {i}")
        break

if linha_dados is None:
    print("ERRO: Linha de dados não encontrada!")
    exit(1)

# Corrigir o mapeamento das notas
row = tabela_notas.rows[linha_dados]

print("\nCorrigindo mapeamento de notas:")
print("Células antes:")
for i, cell in enumerate(row.cells):
    print(f"  Célula {i}: {cell.text[:50]}")

# Mapear corretamente as células
# Baseado no debug anterior:
# Célula 0: {{loop.index}}
# Célula 1: {{aluno.nome}}
# Células 2-11: notas (av1, av2, av3, av4, av5) - cada uma aparece 2x por causa de merge
# Célula 12-13: {{aluno.media}}

# Corrigir as células de notas
mapeamento = {
    0: '{{loop.index}}',
    1: '{{aluno.nome}}',
    2: '{{aluno.av1}}',
    3: '{{aluno.av1}}',  # Mesclada
    4: '{{aluno.av2}}',
    5: '{{aluno.av2}}',  # Mesclada
    6: '{{aluno.av3}}',
    7: '{{aluno.av3}}',  # Mesclada
    8: '{{aluno.av4}}',
    9: '{{aluno.av4}}',  # Mesclada
    10: '{{aluno.av5}}',
    11: '{{aluno.av5}}',  # Mesclada
    12: '{{aluno.media}}',
    13: '{{aluno.media}}',  # Mesclada
}

for i, cell in enumerate(row.cells):
    if i in mapeamento:
        cell.text = ''
        cell.paragraphs[0].text = mapeamento[i]
        print(f"  ✓ Célula {i} atualizada: {mapeamento[i]}")

# Salvar
output_path = 'templates_quadros/notas/quadro_notas_template.docx'
doc.save(output_path)

print(f"\n✅ Template corrigido!")
print("Mapeamento de notas ajustado corretamente.")
