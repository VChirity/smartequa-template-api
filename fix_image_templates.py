from docx import Document
import os

def fix_template_tags(filepath):
    """
    Corrige tags fragmentadas em templates Word para evitar corrupção
    """
    if not os.path.exists(filepath):
        print(f"❌ Arquivo não encontrado: {filepath}")
        return False
    
    print(f"\n🔧 Processando: {filepath}")
    
    try:
        doc = Document(filepath)
        changes_made = 0
        
        # Lista de tags que podem estar fragmentadas
        tags_to_check = [
            'responsavel1', 'cpf_responsavel', 'endereco_completo', 'bairro', 'cep',
            'naturalidade_resp1', 'nasc_resp1',
            'responsavel2', 'cpf2', 'endereco2', 'bairro2', 'cep2',
            'nome_aluno', 'ano', 'ano_letivo', 'data_extenso',
            'naturalidade_aluno', 'nasc_aluno', 'cpf_aluno',
            'desconto', 'desconto_extenso'
        ]
        
        # Processa todos os parágrafos
        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text
            
            # Verifica se há tags no parágrafo
            if '{{' in text and '}}' in text:
                # Reconstrói o parágrafo consolidando runs
                full_text = para.text
                
                # Limpa todos os runs
                for run in para.runs:
                    run.text = ''
                
                # Adiciona o texto completo em um único run
                if para.runs:
                    para.runs[0].text = full_text
                    changes_made += 1
                else:
                    para.add_run(full_text)
                    changes_made += 1
        
        # Processa tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = para.text
                        
                        if '{{' in text and '}}' in text:
                            full_text = para.text
                            
                            for run in para.runs:
                                run.text = ''
                            
                            if para.runs:
                                para.runs[0].text = full_text
                                changes_made += 1
                            else:
                                para.add_run(full_text)
                                changes_made += 1
        
        if changes_made > 0:
            # Salva o arquivo corrigido
            doc.save(filepath)
            print(f"✅ Corrigido! {changes_made} parágrafos consolidados")
            return True
        else:
            print(f"ℹ️ Nenhuma tag fragmentada encontrada")
            return True
            
    except Exception as e:
        print(f"❌ Erro ao processar: {e}")
        import traceback
        traceback.print_exc()
        return False

def main():
    print("=" * 70)
    print("FIX IMAGE TEMPLATES - Correção de Tags Fragmentadas")
    print("=" * 70)
    
    templates = [
        "templates/IMAGEM-E-VOZ-ALUNO-PUBLICIDADE_1.docx",
        "templates/IMAGEM-E-VOZ-ALUNO-INSTITUCIONAL_1.docx"
    ]
    
    success_count = 0
    for template_path in templates:
        if fix_template_tags(template_path):
            success_count += 1
    
    print("\n" + "=" * 70)
    print(f"✅ Concluído! {success_count}/{len(templates)} arquivos processados")
    print("=" * 70)
    
    if success_count == len(templates):
        print("\n🎉 Todos os templates estão prontos para uso!")
    else:
        print("\n⚠️ Alguns templates podem ter problemas. Verifique os erros acima.")

if __name__ == '__main__':
    main()
