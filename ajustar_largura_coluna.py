from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Carregar o template corrigido
template_path = 'templates_quadros/notas/quadro_notas_template_CORRIGIDO.docx'

print(f"Carregando template: {template_path}")
doc = Document(template_path)

# Encontrar a tabela de notas
tabela_notas = None
for table in doc.tables:
    if len(table.columns) >= 8:
        tabela_notas = table
        break

if not tabela_notas:
    print("ERRO: Tabela não encontrada!")
    exit(1)

print(f"Tabela encontrada com {len(tabela_notas.columns)} colunas")

# Ajustar largura da primeira coluna (Nº) - tornar mais estreita
# Largura atual é muito grande, vamos reduzir para 1.5cm
tabela_notas.columns[0].width = Cm(1.5)
print("✓ Largura da coluna Nº ajustada para 1.5cm")

# Centralizar conteúdo da primeira coluna em todas as linhas
for i, row in enumerate(tabela_notas.rows):
    cell = row.cells[0]
    # Centralizar todos os parágrafos da célula
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    print(f"✓ Linha {i}: conteúdo centralizado")

# Salvar
output_path = 'templates_quadros/notas/quadro_notas_template_FINAL.docx'
doc.save(output_path)

print(f"\n✅ Template ajustado salvo em: {output_path}")
print("Agora teste novamente!")
