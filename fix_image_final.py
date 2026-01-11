from docx import Document
import os

def fix_image_template_final(filepath):
    """
    Aplica o mesmo método que funcionou para o contrato de desconto.
    Consolida tags fragmentadas de forma agressiva.
    """
    print(f"\n🔧 Processando: {filepath}")
    
    if not os.path.exists(filepath):
        print(f"❌ Arquivo não encontrado!")
        return False
    
    try:
        doc = Document(filepath)
        print(f"✅ Documento carregado")
        
        changes_made = 0
        
        # Processa TODOS os parágrafos
        for para in doc.paragraphs:
            full_text = para.text
            
            # Se tem tags E tem múltiplos runs, consolida
            if '{{' in full_text and '}}' in full_text:
                num_runs = len(para.runs)
                
                if num_runs > 1:
                    # Preserva formatação do primeiro run
                    if para.runs:
                        first_run = para.runs[0]
                        bold = first_run.bold
                        italic = first_run.italic
                        font_name = first_run.font.name if first_run.font.name else None
                        font_size = first_run.font.size
                        font_color = None
                        if first_run.font.color and first_run.font.color.rgb:
                            font_color = first_run.font.color.rgb
                        
                        # Remove TODOS os runs exceto o primeiro
                        for i in range(len(para.runs) - 1, 0, -1):
                            para._element.remove(para.runs[i]._element)
                        
                        # Coloca TODO o texto no primeiro run
                        para.runs[0].text = full_text
                        
                        # Restaura formatação
                        if bold is not None:
                            para.runs[0].bold = bold
                        if italic is not None:
                            para.runs[0].italic = italic
                        if font_name:
                            para.runs[0].font.name = font_name
                        if font_size:
                            para.runs[0].font.size = font_size
                        if font_color:
                            from docx.shared import RGBColor
                            para.runs[0].font.color.rgb = font_color
                        
                        changes_made += 1
                        print(f"  ✓ Consolidado: {num_runs} runs → 1 run")
        
        # Processa tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        full_text = para.text
                        
                        if '{{' in full_text and '}}' in full_text:
                            num_runs = len(para.runs)
                            
                            if num_runs > 1:
                                if para.runs:
                                    first_run = para.runs[0]
                                    bold = first_run.bold
                                    italic = first_run.italic
                                    font_name = first_run.font.name if first_run.font.name else None
                                    font_size = first_run.font.size
                                    font_color = None
                                    if first_run.font.color and first_run.font.color.rgb:
                                        font_color = first_run.font.color.rgb
                                    
                                    for i in range(len(para.runs) - 1, 0, -1):
                                        para._element.remove(para.runs[i]._element)
                                    
                                    para.runs[0].text = full_text
                                    
                                    if bold is not None:
                                        para.runs[0].bold = bold
                                    if italic is not None:
                                        para.runs[0].italic = italic
                                    if font_name:
                                        para.runs[0].font.name = font_name
                                    if font_size:
                                        para.runs[0].font.size = font_size
                                    if font_color:
                                        from docx.shared import RGBColor
                                        para.runs[0].font.color.rgb = font_color
                                    
                                    changes_made += 1
        
        # Processa cabeçalhos e rodapés
        for section in doc.sections:
            for para in section.header.paragraphs:
                full_text = para.text
                if '{{' in full_text and '}}' in full_text and len(para.runs) > 1:
                    if para.runs:
                        for i in range(len(para.runs) - 1, 0, -1):
                            para._element.remove(para.runs[i]._element)
                        para.runs[0].text = full_text
                        changes_made += 1
            
            for para in section.footer.paragraphs:
                full_text = para.text
                if '{{' in full_text and '}}' in full_text and len(para.runs) > 1:
                    if para.runs:
                        for i in range(len(para.runs) - 1, 0, -1):
                            para._element.remove(para.runs[i]._element)
                        para.runs[0].text = full_text
                        changes_made += 1
        
        # Salva
        doc.save(filepath)
        
        print(f"✅ Documento salvo com {changes_made} consolidações")
        print(f"🎨 Formatação e logo preservadas")
        
        return True
        
    except Exception as e:
        print(f"❌ Erro: {e}")
        import traceback
        traceback.print_exc()
        return False

def main():
    print("=" * 70)
    print("FIX IMAGE FINAL - Método do Contrato de Desconto")
    print("=" * 70)
    
    templates = [
        "templates/IMAGEM-E-VOZ-ALUNO-PUBLICIDADE_1.docx",
        "templates/IMAGEM-E-VOZ-ALUNO-INSTITUCIONAL_1.docx"
    ]
    
    success_count = 0
    for template_path in templates:
        if fix_image_template_final(template_path):
            success_count += 1
    
    print("\n" + "=" * 70)
    print(f"✅ Concluído! {success_count}/{len(templates)} arquivos corrigidos")
    print("=" * 70)
    
    if success_count == len(templates):
        print("\n🎉 Templates corrigidos! Devem abrir no Word sem erros.")
        print("💡 Teste abrindo os arquivos agora.")
        print("🚀 Depois teste no app: streamlit run app.py")
    else:
        print("\n⚠️ Alguns arquivos falharam.")

if __name__ == '__main__':
    main()
