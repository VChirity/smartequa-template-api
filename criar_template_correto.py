from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_background(cell, color):
    """Definir cor de fundo da célula"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

# Criar documento
doc = Document()

# Configurar margens
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

# Título
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('QUADRO DE NOTAS')
run.bold = True
run.font.size = Pt(16)

# Subtítulo com tags
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('{{bimestre}} — {{ano}}')
run.font.size = Pt(12)

# Espaço
doc.add_paragraph()

# Turma
turma_para = doc.add_paragraph()
turma_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = turma_para.add_run('{{turma}}')
run.bold = True
run.font.size = Pt(14)

# Espaço
doc.add_paragraph()

# Tabela de notas - APENAS CABEÇALHO
table = doc.add_table(rows=2, cols=8)
table.style = 'Table Grid'

# Linha 0: Cabeçalho "NOTAS"
row0 = table.rows[0]
cell = row0.cells[0]
for i in range(1, 8):
    cell.merge(row0.cells[i])
cell.text = 'NOTAS'
cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
run = cell.paragraphs[0].runs[0]
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(255, 255, 255)
set_cell_background(cell, '1F4E78')

# Linha 1: Cabeçalhos das colunas
headers = ['Nº', 'NOME DO ALUNO', '{{tipo_av_1}}\n{{pont_1}}', '{{tipo_av_2}}\n{{pont_2}}', 
           '{{tipo_av_3}}\n{{pont_3}}', '{{tipo_av_4}}\n{{pont_4}}', '{{tipo_av_5}}\n{{pont_5}}', 'MÉDIA']
row1 = table.rows[1]
for i, header in enumerate(headers):
    cell = row1.cells[i]
    cell.text = header
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cell.paragraphs[0].runs[0]
    run.bold = True
    run.font.size = Pt(10)
    if i < 2:
        set_cell_background(cell, '5B9BD5')
    elif i == 7:
        set_cell_background(cell, '70AD47')
    else:
        set_cell_background(cell, 'FFC000')

# ADICIONAR LINHAS DO LOOP DINAMICAMENTE
# Linha para início do loop
row_loop_start = table.add_row()
cell = row_loop_start.cells[0]
for i in range(1, 8):
    cell.merge(row_loop_start.cells[i])
# Adicionar o texto do loop diretamente no XML
para = cell.paragraphs[0]
para.text = '{%tr for aluno in alunos %}'
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Linha de dados
row_data = table.add_row()
dados = ['{{loop.index}}', '{{aluno.nome}}', '{{aluno.av1}}', '{{aluno.av2}}', 
         '{{aluno.av3}}', '{{aluno.av4}}', '{{aluno.av5}}', '{{aluno.media}}']
for i, dado in enumerate(dados):
    cell = row_data.cells[i]
    cell.text = dado
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Linha para fim do loop
row_loop_end = table.add_row()
cell = row_loop_end.cells[0]
for i in range(1, 8):
    cell.merge(row_loop_end.cells[i])
para = cell.paragraphs[0]
para.text = '{%tr endfor %}'
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Espaço
doc.add_paragraph()
doc.add_paragraph()

# Cálculo da média
doc.add_paragraph('CÁLCULO DA MÉDIA')
calculo_para = doc.add_paragraph()
calculo_para.add_run('☐ Nota única    ☐ Média Aritmética    ☐ Adição das notas    ☐ Outro')
doc.add_paragraph('Como calcular a média: {{explicacao_calculo}}')

# Espaço
doc.add_paragraph()

# Rodapé
footer_para = doc.add_paragraph()
footer_para.add_run('Professor: {{professor}}')
footer_para = doc.add_paragraph()
footer_para.add_run('Disciplina: {{disciplina}}')

# Salvar
output_path = 'templates_quadros/notas/quadro_notas_template_CORRETO.docx'
doc.save(output_path)
print(f'✅ Template criado: {output_path}')
print('\nEste template usa {%tr %} para loops em tabelas.')
print('Teste com o script test_template.py')
