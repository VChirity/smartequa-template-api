from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def copy_to_fresh_document(source_path, dest_path):
    """
    Cria um documento Word completamente novo e copia o conteúdo do arquivo corrompido.
    Isso gera uma estrutura XML limpa que o Word aceita.
    """
    print(f"\n🔧 Processando: {os.path.basename(dest_path)}")
    
    if not os.path.exists(source_path):
        print(f"❌ Arquivo fonte não encontrado!")
        return False
    
    try:
        # Carrega o arquivo corrompido (que o python-docx consegue ler)
        print("📖 Lendo arquivo corrompido...")
        source_doc = Document(source_path)
        
        # Cria um documento NOVO e VAZIO
        print("📄 Criando documento novo e vazio...")
        new_doc = Document()
        
        # Configura página igual ao original
        section = new_doc.sections[0]
        section.page_height = Inches(11.69)  # A4
        section.page_width = Inches(8.27)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        
        # Copia cada parágrafo do original para o novo
        print("✏️ Copiando conteúdo parágrafo por parágrafo...")
        for para in source_doc.paragraphs:
            # Cria novo parágrafo
            new_para = new_doc.add_paragraph()
            
            # Copia cada run do parágrafo original
            for run in para.runs:
                new_run = new_para.add_run(run.text)
                
                # Copia formatação
                if run.bold is not None:
                    new_run.bold = run.bold
                if run.italic is not None:
                    new_run.italic = run.italic
                if run.underline is not None:
                    new_run.underline = run.underline
                
                # Copia fonte
                if run.font.name:
                    new_run.font.name = run.font.name
                if run.font.size:
                    new_run.font.size = run.font.size
                
                # Copia cor
                if run.font.color and run.font.color.rgb:
                    new_run.font.color.rgb = run.font.color.rgb
            
            # Copia alinhamento do parágrafo
            if para.alignment:
                new_para.alignment = para.alignment
        
        # Copia tabelas se houver
        if source_doc.tables:
            print("📊 Copiando tabelas...")
            for table in source_doc.tables:
                # Cria nova tabela
                new_table = new_doc.add_table(rows=len(table.rows), cols=len(table.columns))
                
                # Copia conteúdo de cada célula
                for i, row in enumerate(table.rows):
                    for j, cell in enumerate(row.cells):
                        new_cell = new_table.rows[i].cells[j]
                        
                        # Limpa célula nova
                        for para in new_cell.paragraphs:
                            para.clear()
                        
                        # Copia parágrafos da célula original
                        for para in cell.paragraphs:
                            new_para = new_cell.add_paragraph()
                            for run in para.runs:
                                new_run = new_para.add_run(run.text)
                                if run.bold is not None:
                                    new_run.bold = run.bold
                                if run.italic is not None:
                                    new_run.italic = run.italic
        
        # Salva o documento novo
        print(f"💾 Salvando documento novo: {dest_path}")
        new_doc.save(dest_path)
        
        print(f"✅ Documento criado com estrutura XML limpa!")
        return True
        
    except Exception as e:
        print(f"❌ Erro: {e}")
        import traceback
        traceback.print_exc()
        return False

def main():
    print("=" * 70)
    print("COPY TO FRESH DOCUMENT - Método Copiar/Colar")
    print("Cria documentos novos e copia conteúdo = XML limpo")
    print("=" * 70)
    
    templates = [
        {
            'source': 'templates backup/IMAGEM-E-VOZ-ALUNO-PUBLICIDADE_1.docx',
            'dest': 'templates/IMAGEM-E-VOZ-ALUNO-PUBLICIDADE_1.docx'
        },
        {
            'source': 'templates backup/IMAGEM-E-VOZ-ALUNO-INSTITUCIONAL_1.docx',
            'dest': 'templates/IMAGEM-E-VOZ-ALUNO-INSTITUCIONAL_1.docx'
        }
    ]
    
    success_count = 0
    for template in templates:
        if copy_to_fresh_document(template['source'], template['dest']):
            success_count += 1
    
    print("\n" + "=" * 70)
    print(f"✅ Concluído! {success_count}/{len(templates)} arquivos criados")
    print("=" * 70)
    
    if success_count == len(templates):
        print("\n🎉 Documentos novos criados com XML limpo!")
        print("💡 Método: Copiar/Colar programático")
        print("🔍 Tente abrir no Word agora - deve funcionar!")
    else:
        print("\n⚠️ Alguns arquivos falharam.")

if __name__ == '__main__':
    main()
