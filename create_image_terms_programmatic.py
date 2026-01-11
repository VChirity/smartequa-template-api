from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_image_term_programmatic(source, dest, title, content_text):
    """
    Cria termo de imagem de forma programática (mesmo método do contrato de desconto).
    Insere tags em runs separados e completos para evitar fragmentação.
    """
    print(f"\n🔧 Criando: {os.path.basename(dest)}")
    print(f"📂 Fonte: {source}")
    
    if not os.path.exists(source):
        print(f"❌ Arquivo fonte não encontrado!")
        return False
    
    try:
        # Carrega o template base
        doc = Document(source)
        print("✅ Template base carregado")
        
        # Remove todos os parágrafos EXCETO o primeiro (logo/cabeçalho)
        print("🗑️ Removendo conteúdo antigo (preservando cabeçalho)...")
        paragraphs_to_remove = list(doc.paragraphs[1:])
        for para in paragraphs_to_remove:
            p = para._element
            p.getparent().remove(p)
        
        # Adiciona título
        print("📝 Adicionando título...")
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(14)
        title_run.font.name = 'Calibri'
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Linha em branco
        doc.add_paragraph()
        
        # Adiciona conteúdo com tags inseridas programaticamente
        print("✏️ Inserindo conteúdo com tags programáticas...")
        
        # Parágrafo 1 - Texto principal
        p1 = doc.add_paragraph()
        p1.add_run('Eu, ')
        r = p1.add_run('{{responsavel1}}')
        r.bold = True
        p1.add_run(', ')
        r = p1.add_run('{{naturalidade_resp1}}')
        r.bold = True
        p1.add_run(', nascido(a) em ')
        r = p1.add_run('{{nasc_resp1}}')
        r.bold = True
        p1.add_run(', inscrito(a) no CPF/MF sob nº ')
        r = p1.add_run('{{cpf_responsavel}}')
        r.bold = True
        p1.add_run(', residente no endereço ')
        r = p1.add_run('{{endereco_completo}}')
        r.bold = True
        p1.add_run(', Rio de Janeiro – RJ, responsável pelo/pela criança/adolescente ')
        r = p1.add_run('{{nome_aluno}}')
        r.bold = True
        p1.add_run(', ')
        r = p1.add_run('{{naturalidade_aluno}}')
        r.bold = True
        p1.add_run(', nascido(a) em ')
        r = p1.add_run('{{nasc_aluno}}')
        r.bold = True
        p1.add_run(', inscrito(a) no CPF/MF sob nº ')
        r = p1.add_run('{{cpf_aluno}}')
        r.bold = True
        p1.add_run(', pelo presente instrumento, AUTORIZO o CURSO DE ESPECIALIZAÇÃO EQUAÇÃO LTDA, com sede na Rua Mendes Tavares, nº108, Vila Isabel, Rio de Janeiro – RJ, inscrita no CNPJ/MF sob o nº 42.319.202.001-40, a fazer uso da imagem e/ou voz do menor acima identificado, em todo e qualquer material entre fotos, documentos e outros meios de comunicação, para ')
        p1.add_run(content_text)  # "campanhas publicitárias" ou "campanhas institucionais"
        p1.add_run(', sejam essas destinadas à divulgação ao público em geral e/ou apenas para uso desta escola.')
        p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Linha em branco
        doc.add_paragraph()
        
        # Parágrafo 2
        p2 = doc.add_paragraph()
        p2.add_run('A presente autorização é concedida a título gratuito, abrangendo o uso da imagem acima mencionada em todo território nacional e no exterior, sob qualquer forma e meios, ou sejam, em destaque: (I) outdoor; (II) busdoor; folhetos em geral (encartes, mala direta, catálogo, etc.); (III) folder de apresentação; (IV) anúncios em revistas e jornais em geral; (V) home page; (VI) cartazes; (VII) backlight; (VIII) mídia eletrônica (internet, painéis, vídeotapes, televisão, cinema, programa para rádio, entre outros).')
        p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Linha em branco
        doc.add_paragraph()
        
        # Parágrafo 3
        p3 = doc.add_paragraph()
        p3.add_run('Por esta ser a expressão da minha vontade, declaro que autorizo o uso acima descrito sem que nada haja a ser reclamado, e assino a presente autorização em 02 (duas) vias de igual teor e forma.')
        p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Linha em branco
        doc.add_paragraph()
        
        # Data
        p4 = doc.add_paragraph()
        p4.add_run('Rio de Janeiro, ')
        r = p4.add_run('{{data_extenso}}')
        r.bold = True
        p4.add_run('.')
        p4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Linhas em branco
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Assinatura
        p5 = doc.add_paragraph()
        p5.add_run('________________________________________________')
        p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p6 = doc.add_paragraph()
        r = p6.add_run('{{responsavel1}}')
        r.bold = True
        p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p7 = doc.add_paragraph()
        p7.add_run('CPF: ')
        r = p7.add_run('{{cpf_responsavel}}')
        r.bold = True
        p7.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Salva o documento
        print(f"💾 Salvando: {dest}")
        doc.save(dest)
        
        print("✅ Termo criado com sucesso!")
        print("🔍 Tags inseridas programaticamente (sem fragmentação)")
        return True
        
    except Exception as e:
        print(f"❌ Erro: {e}")
        import traceback
        traceback.print_exc()
        return False

def main():
    print("=" * 70)
    print("CREATE IMAGE TERMS - Método Programático (como desconto)")
    print("=" * 70)
    
    source = 'templates/template_contrato2025_2.docx'
    
    templates = [
        {
            'dest': 'templates/IMAGEM-E-VOZ-ALUNO-PUBLICIDADE_1.docx',
            'title': 'TERMO DE AUTORIZAÇÃO DE USO\nIMAGEM E VOZ DE ALUNO - PUBLICIDADE',
            'content': 'campanhas publicitárias'
        },
        {
            'dest': 'templates/IMAGEM-E-VOZ-ALUNO-INSTITUCIONAL_1.docx',
            'title': 'TERMO DE AUTORIZAÇÃO DE USO\nIMAGEM E VOZ DE ALUNO - INSTITUCIONAL',
            'content': 'campanhas institucionais'
        }
    ]
    
    success_count = 0
    for template in templates:
        if create_image_term_programmatic(source, template['dest'], template['title'], template['content']):
            success_count += 1
    
    print("\n" + "=" * 70)
    print(f"✅ Concluído! {success_count}/{len(templates)} termos criados")
    print("=" * 70)
    
    if success_count == len(templates):
        print("\n🎉 Termos criados com MESMO MÉTODO do contrato de desconto!")
        print("💡 Tags inseridas programaticamente = SEM fragmentação")
        print("🧪 Teste no app agora - deve funcionar!")
    else:
        print("\n⚠️ Alguns termos falharam.")

if __name__ == '__main__':
    main()
