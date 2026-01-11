from docx import Document
from docx.shared import RGBColor
import os

def fix_all_red_text(filepath):
    """
    Substitui TODOS os textos em vermelho pelas tags corretas em preto
    """
    print(f"\n🔧 Processando: {filepath}")
    
    if not os.path.exists(filepath):
        print(f"❌ Arquivo não encontrado!")
        return False
    
    try:
        doc = Document(filepath)
        print(f"✅ Documento carregado")
        
        # Mapeamento COMPLETO baseado na imagem
        replacements = {
            '(NATURALIDADE DO RESPONSÁVEL)': '{{naturalidade_resp1}}',
            '(DATA DE NASCIMENTO DO RESPONSÁVEL)': '{{nasc_resp1}}',
            '(CPF DO RESPONSÁVEL)': '{{cpf_responsavel}}',
            '(ENDEREÇO COMPLETO DO RESPONSÁVEL)': '{{endereco_completo}}',
            '{{endereco_completo}}': '{{endereco_completo}}',  # Já correto
            '{{nome_aluno}}': '{{nome_aluno}}',  # Já correto
            '(NATURALIDADE DO ALUNO)': '{{naturalidade_aluno}}',
            '(DATA DE NASCIMENTO DO ALUNO)': '{{nasc_aluno}}',
            '(CPF DO ALUNO)': '{{cpf_aluno}}',
            '(DATA POR EXTENSO)': '{{data_extenso}}',
            'DATA POR EXTENSO': '{{data_extenso}}',
        }
        
        changes_made = 0
        red_text_found = 0
        
        # Processa todos os parágrafos
        for para in doc.paragraphs:
            for run in para.runs:
                # Verifica se o texto está em vermelho
                is_red = False
                if run.font.color and run.font.color.rgb:
                    rgb = run.font.color.rgb
                    # Vermelho é RGB(255, 0, 0) ou próximo
                    if rgb[0] > 200 and rgb[1] < 100 and rgb[2] < 100:
                        is_red = True
                        red_text_found += 1
                
                original_text = run.text
                new_text = original_text
                
                # Substitui cada placeholder
                for placeholder, tag in replacements.items():
                    if placeholder in new_text:
                        new_text = new_text.replace(placeholder, tag)
                        changes_made += 1
                        print(f"  ✓ Substituído: {placeholder} → {tag}")
                
                # Se houve mudança OU se estava em vermelho, atualiza
                if new_text != original_text or is_red:
                    run.text = new_text
                    # Define cor como preta (remove qualquer cor customizada)
                    run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Processa tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            is_red = False
                            if run.font.color and run.font.color.rgb:
                                rgb = run.font.color.rgb
                                if rgb[0] > 200 and rgb[1] < 100 and rgb[2] < 100:
                                    is_red = True
                                    red_text_found += 1
                            
                            original_text = run.text
                            new_text = original_text
                            
                            for placeholder, tag in replacements.items():
                                if placeholder in new_text:
                                    new_text = new_text.replace(placeholder, tag)
                                    changes_made += 1
                                    print(f"  ✓ Substituído: {placeholder} → {tag}")
                            
                            if new_text != original_text or is_red:
                                run.text = new_text
                                run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Processa cabeçalhos e rodapés
        for section in doc.sections:
            for para in section.header.paragraphs:
                for run in para.runs:
                    is_red = False
                    if run.font.color and run.font.color.rgb:
                        rgb = run.font.color.rgb
                        if rgb[0] > 200 and rgb[1] < 100 and rgb[2] < 100:
                            is_red = True
                            red_text_found += 1
                    
                    original_text = run.text
                    new_text = original_text
                    
                    for placeholder, tag in replacements.items():
                        if placeholder in new_text:
                            new_text = new_text.replace(placeholder, tag)
                            changes_made += 1
                            print(f"  ✓ Substituído: {placeholder} → {tag}")
                    
                    if new_text != original_text or is_red:
                        run.text = new_text
                        run.font.color.rgb = RGBColor(0, 0, 0)
            
            for para in section.footer.paragraphs:
                for run in para.runs:
                    is_red = False
                    if run.font.color and run.font.color.rgb:
                        rgb = run.font.color.rgb
                        if rgb[0] > 200 and rgb[1] < 100 and rgb[2] < 100:
                            is_red = True
                            red_text_found += 1
                    
                    original_text = run.text
                    new_text = original_text
                    
                    for placeholder, tag in replacements.items():
                        if placeholder in new_text:
                            new_text = new_text.replace(placeholder, tag)
                            changes_made += 1
                            print(f"  ✓ Substituído: {placeholder} → {tag}")
                    
                    if new_text != original_text or is_red:
                        run.text = new_text
                        run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Salva o documento
        doc.save(filepath)
        
        print(f"✅ Documento salvo")
        print(f"📊 {changes_made} substituições realizadas")
        print(f"🎨 {red_text_found} textos em vermelho convertidos para preto")
        
        return True
        
    except Exception as e:
        print(f"❌ Erro: {e}")
        import traceback
        traceback.print_exc()
        return False

def main():
    print("=" * 70)
    print("FIX ALL RED TEXT - Substituição Completa")
    print("Remove TODO texto vermelho e insere tags em PRETO")
    print("=" * 70)
    
    templates = [
        "templates/IMAGEM-E-VOZ-ALUNO-PUBLICIDADE_1.docx",
        "templates/IMAGEM-E-VOZ-ALUNO-INSTITUCIONAL_1.docx"
    ]
    
    success_count = 0
    for template_path in templates:
        if fix_all_red_text(template_path):
            success_count += 1
    
    print("\n" + "=" * 70)
    print(f"✅ Concluído! {success_count}/{len(templates)} arquivos processados")
    print("=" * 70)
    
    if success_count == len(templates):
        print("\n🎉 TODO texto vermelho removido e substituído por tags em PRETO!")
        print("💡 Abra os arquivos no Word para verificar.")
        print("📝 Não deve haver NADA em vermelho.")
    else:
        print("\n⚠️ Alguns arquivos falharam.")

if __name__ == '__main__':
    main()
