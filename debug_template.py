from docx import Document

template_path = 'templates_quadros/notas/quadro_notas_template.docx'
doc = Document(template_path)

print("=== PARÁGRAFOS ===")
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        print(f"{i}: {para.text}")

print("\n=== TABELAS ===")
for t_idx, table in enumerate(doc.tables):
    print(f"\nTabela {t_idx}:")
    for r_idx, row in enumerate(table.rows):
        print(f"  Linha {r_idx}:")
        for c_idx, cell in enumerate(row.cells):
            if cell.text.strip():
                print(f"    Célula {c_idx}: {cell.text[:100]}")
