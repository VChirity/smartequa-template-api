from docx import Document
from docxtpl import DocxTemplate
import os

def verify_template(filepath, template_name):
    """
    Verifica se o template está funcionando corretamente
    """
    print(f"\n{'='*70}")
    print(f"Verificando: {template_name}")
    print(f"{'='*70}")
    
    if not os.path.exists(filepath):
        print(f"❌ Arquivo não encontrado!")
        return False
    
    try:
        # Teste 1: Abre com python-docx
        print("\n1️⃣ Teste python-docx:")
        doc = Document(filepath)
        print(f"   ✅ Arquivo carrega OK")
        print(f"   📊 Parágrafos: {len(doc.paragraphs)}")
        print(f"   📊 Tabelas: {len(doc.tables)}")
        
        # Teste 2: Verifica tags
        print("\n2️⃣ Verificando tags:")
        tags_found = []
        for para in doc.paragraphs:
            text = para.text
            if '{{' in text and '}}' in text:
                # Extrai tags
                import re
                found = re.findall(r'\{\{([^}]+)\}\}', text)
                tags_found.extend(found)
        
        if tags_found:
            print(f"   ✅ {len(tags_found)} tags encontradas:")
            for tag in set(tags_found):
                print(f"      • {{{{tag}}}}")
        else:
            print(f"   ⚠️ Nenhuma tag encontrada")
        
        # Teste 3: Verifica fragmentação
        print("\n3️⃣ Verificando fragmentação:")
        fragmented = 0
        for para in doc.paragraphs:
            if '{{' in para.text and '}}' in para.text:
                if len(para.runs) > 1:
                    fragmented += 1
        
        if fragmented == 0:
            print(f"   ✅ Nenhum parágrafo com tags fragmentadas")
        else:
            print(f"   ⚠️ {fragmented} parágrafos ainda fragmentados")
        
        # Teste 4: Tenta renderizar com docxtpl
        print("\n4️⃣ Teste docxtpl:")
        try:
            tpl = DocxTemplate(filepath)
            
            # Context de teste
            test_context = {
                'responsavel1': 'João da Silva',
                'cpf_responsavel': '123.456.789-00',
                'endereco_completo': 'Rua Teste, 123',
                'bairro': 'Centro',
                'cep': '20000-000',
                'naturalidade_resp1': 'Rio de Janeiro',
                'nasc_resp1': '01/01/1980',
                'responsavel2': 'Maria Silva',
                'cpf2': '987.654.321-00',
                'endereco2': 'Rua Teste 2, 456',
                'bairro2': 'Zona Sul',
                'cep2': '20001-000',
                'nome_aluno': 'Pedro Silva',
                'ano': '5º Ano',
                'ano_letivo': '2025',
                'data_extenso': '10 de janeiro de 2026',
                'naturalidade_aluno': 'Rio de Janeiro',
                'nasc_aluno': '15/03/2010',
                'cpf_aluno': '111.222.333-44',
                'desconto': '10',
                'desconto_extenso': 'dez'
            }
            
            tpl.render(test_context)
            print(f"   ✅ Renderização OK com docxtpl")
            
            # Tenta salvar em memória
            import io
            bio = io.BytesIO()
            tpl.save(bio)
            print(f"   ✅ Salvamento em memória OK")
            print(f"   📊 Tamanho do arquivo gerado: {len(bio.getvalue())} bytes")
            
        except Exception as e:
            print(f"   ❌ Erro no docxtpl: {e}")
            return False
        
        print(f"\n{'='*70}")
        print(f"✅ {template_name}: TUDO OK!")
        print(f"{'='*70}")
        return True
        
    except Exception as e:
        print(f"\n❌ ERRO: {e}")
        import traceback
        traceback.print_exc()
        return False

def main():
    print("\n" + "=" * 70)
    print("VERIFICAÇÃO COMPLETA DE TEMPLATES")
    print("=" * 70)
    
    templates = [
        ("templates/template_contrato2025_2.docx", "Contrato Padrão"),
        ("templates/template_contratoDESCONTO2025_2.docx", "Contrato com Desconto"),
        ("templates/IMAGEM-E-VOZ-ALUNO-PUBLICIDADE_1.docx", "Termo Imagem Publicidade"),
        ("templates/IMAGEM-E-VOZ-ALUNO-INSTITUCIONAL_1.docx", "Termo Imagem Institucional")
    ]
    
    results = {}
    for filepath, name in templates:
        results[name] = verify_template(filepath, name)
    
    print("\n" + "=" * 70)
    print("RESUMO FINAL")
    print("=" * 70)
    
    for name, success in results.items():
        status = "✅ OK" if success else "❌ FALHOU"
        print(f"{status} - {name}")
    
    all_ok = all(results.values())
    
    if all_ok:
        print("\n🎉 TODOS OS TEMPLATES ESTÃO FUNCIONANDO!")
        print("💡 Pode testar no Word e no app sem problemas.")
    else:
        print("\n⚠️ Alguns templates têm problemas.")
        print("💡 Verifique os detalhes acima.")

if __name__ == '__main__':
    main()
