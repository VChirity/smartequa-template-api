from docx import Document
from docx.shared import RGBColor
import os

def final_tag_replacement(filepath):
    """
    Substitui TODOS os placeholders com parênteses pelas tags corretas
    Mapeamento completo fornecido pelo usuário
    """
    print(f"\n🔧 Processando: {filepath}")
    
    if not os.path.exists(filepath):
        print(f"❌ Arquivo não encontrado!")
        return False
    
    try:
        doc = Document(filepath)
        print(f"✅ Documento carregado")
        
        # Mapeamento COMPLETO fornecido pelo usuário
        replacements = {
            '(NOME COMPLETO DO RESPONSÁVEL)': '{{responsavel1}}',
            '(CPF DO RESPONSÁVEL)': '{{cpf_responsavel}}',
            '(ENDEREÇO COMPLETO DO RESPONSÁVEL)': '{{endereco_completo}}',
            '(NATURALIDADE DO RESPONSÁVEL)': '{{naturalidade_resp1}}',
            '(DATA DE NASCIMENTO DO RESPONSÁVEL)': '{{nasc_resp1}}',
            '(NOME COMPLETO DO ALUNO)': '{{nome_aluno}}',
            '(NATURALIDADE DO ALUNO)': '{{naturalidade_aluno}}',
            '(DATA DE NASCIMENTO DO ALUNO)': '{{nasc_aluno}}',
            '(CPF DO ALUNO)': '{{cpf_aluno}}',
            '(DATA DO DIA)': '{{data_extenso}}',
            # Variações sem parênteses também
            'NOME COMPLETO DO RESPONSÁVEL': '{{responsavel1}}',
            'CPF DO RESPONSÁVEL': '{{cpf_responsavel}}',
            'ENDEREÇO COMPLETO DO RESPONSÁVEL': '{{endereco_completo}}',
            'NATURALIDADE DO RESPONSÁVEL': '{{naturalidade_resp1}}',
            'DATA DE NASCIMENTO DO RESPONSÁVEL': '{{nasc_resp1}}',
            'NOME COMPLETO DO ALUNO': '{{nome_aluno}}',
            'NATURALIDADE DO ALUNO': '{{naturalidade_aluno}}',
            'DATA DE NASCIMENTO DO ALUNO': '{{nasc_aluno}}',
            'CPF DO ALUNO': '{{cpf_aluno}}',
            'DATA DO DIA': '{{data_extenso}}',
        }
        
        changes_made = 0
        
        # Processa todos os parágrafos
        for para in doc.paragraphs:
            for run in para.runs:
                original_text = run.text
                new_text = original_text
                
                # Substitui cada placeholder
                for placeholder, tag in replacements.items():
                    if placeholder in new_text:
                        new_text = new_text.replace(placeholder, tag)
                        changes_made += 1
                        print(f"  ✓ Substituído: {placeholder} → {tag}")
                
                # Se houve mudança, atualiza e garante cor preta
                if new_text != original_text:
                    run.text = new_text
                    run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Processa tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            original_text = run.text
                            new_text = original_text
                            
                            for placeholder, tag in replacements.items():
                                if placeholder in new_text:
                                    new_text = new_text.replace(placeholder, tag)
                                    changes_made += 1
                                    print(f"  ✓ Substituído: {placeholder} → {tag}")
                            
                            if new_text != original_text:
                                run.text = new_text
                                run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Processa cabeçalhos e rodapés
        for section in doc.sections:
            for para in section.header.paragraphs:
                for run in para.runs:
                    original_text = run.text
                    new_text = original_text
                    
                    for placeholder, tag in replacements.items():
                        if placeholder in new_text:
                            new_text = new_text.replace(placeholder, tag)
                            changes_made += 1
                            print(f"  ✓ Substituído: {placeholder} → {tag}")
                    
                    if new_text != original_text:
                        run.text = new_text
                        run.font.color.rgb = RGBColor(0, 0, 0)
            
            for para in section.footer.paragraphs:
                for run in para.runs:
                    original_text = run.text
                    new_text = original_text
                    
                    for placeholder, tag in replacements.items():
                        if placeholder in new_text:
                            new_text = new_text.replace(placeholder, tag)
                            changes_made += 1
                            print(f"  ✓ Substituído: {placeholder} → {tag}")
                    
                    if new_text != original_text:
                        run.text = new_text
                        run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Salva o documento
        doc.save(filepath)
        
        print(f"✅ Documento salvo com {changes_made} substituições")
        print(f"🎨 Todas as tags em PRETO, sem parênteses")
        
        return True
        
    except Exception as e:
        print(f"❌ Erro: {e}")
        import traceback
        traceback.print_exc()
        return False

def main():
    print("=" * 70)
    print("FINAL TAG REPLACEMENT - Substituição Definitiva")
    print("Remove TODOS os parênteses e insere tags corretas")
    print("=" * 70)
    
    templates = [
        "templates/IMAGEM-E-VOZ-ALUNO-PUBLICIDADE_1.docx",
        "templates/IMAGEM-E-VOZ-ALUNO-INSTITUCIONAL_1.docx"
    ]
    
    success_count = 0
    for template_path in templates:
        if final_tag_replacement(template_path):
            success_count += 1
    
    print("\n" + "=" * 70)
    print(f"✅ Concluído! {success_count}/{len(templates)} arquivos processados")
    print("=" * 70)
    
    if success_count == len(templates):
        print("\n🎉 TODOS os placeholders substituídos por tags!")
        print("💡 Abra os arquivos no Word para verificar.")
        print("📝 Não deve haver NENHUM parêntese com placeholder.")
    else:
        print("\n⚠️ Alguns arquivos falharam.")

if __name__ == '__main__':
    main()
