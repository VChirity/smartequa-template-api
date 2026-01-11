from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_clean_discount_template():
    """
    Cria template de desconto de forma programática para evitar corrupção de tags
    """
    source = 'templates/template_contrato2025_2.docx'
    dest = 'templates/template_contratoDESCONTO2025_2.docx'
    
    print("🔧 Criando template de desconto limpo...")
    print(f"📂 Fonte: {source}")
    
    if not os.path.exists(source):
        print(f"❌ Arquivo fonte não encontrado!")
        return False
    
    try:
        # Carrega o template base
        doc = Document(source)
        print("✅ Template base carregado")
        
        # Localiza a CLÁUSULA 3ª
        clausula_3_index = None
        for i, para in enumerate(doc.paragraphs):
            if 'CLÁUSULA 3' in para.text.upper() and 'ª' in para.text:
                clausula_3_index = i
                print(f"✅ CLÁUSULA 3ª encontrada no índice {i}")
                break
        
        if clausula_3_index is None:
            print("❌ CLÁUSULA 3ª não encontrada!")
            return False
        
        # Primeiro, renumera as cláusulas de trás para frente
        print("🔄 Renumerando cláusulas...")
        for i in range(len(doc.paragraphs) - 1, clausula_3_index - 1, -1):
            para = doc.paragraphs[i]
            text = para.text
            
            # Substitui números de cláusulas (do maior para o menor)
            for num in range(20, 2, -1):  # De 20 até 3
                if f'CLÁUSULA {num}ª' in text.upper():
                    # Preserva formatação original
                    for run in para.runs:
                        if f'{num}ª' in run.text or f'{num}º' in run.text:
                            run.text = run.text.replace(f'{num}ª', f'{num+1}ª')
                            run.text = run.text.replace(f'{num}º', f'{num+1}º')
                    print(f"  ✓ Renumerado: {num}ª → {num+1}ª")
                    break
        
        # Agora insere a nova CLÁUSULA 3ª
        print("📝 Inserindo nova CLÁUSULA 3ª de desconto...")
        
        # Pega o parágrafo onde estava a CLÁUSULA 3ª (agora 4ª)
        ref_para = doc.paragraphs[clausula_3_index]
        
        # Insere novo parágrafo ANTES
        new_para = ref_para.insert_paragraph_before()
        
        # Adiciona o texto da cláusula em um único run para evitar fragmentação
        run = new_para.add_run(
            'CLÁUSULA 3ª - A contratada concede por mera liberalidade e para o ano letivo vigente, '
            'exclusivamente, um desconto no valor da anuidade devida no percentual de '
        )
        
        # Adiciona tag desconto (em um run separado mas completo)
        run_tag1 = new_para.add_run('{{desconto}}')
        run_tag1.bold = True
        
        run2 = new_para.add_run('% (')
        
        # Adiciona tag desconto_extenso
        run_tag2 = new_para.add_run('{{desconto_extenso}}')
        run_tag2.bold = True
        
        run3 = new_para.add_run(
            '), desde que a mesma seja quitada na data de seu respectivo vencimento, '
            'não implicando tal desconto em novação da anualidade devida.'
        )
        
        # Copia estilo do parágrafo de referência
        if ref_para.style:
            new_para.style = ref_para.style
        
        # Adiciona parágrafo único
        para_unico = ref_para.insert_paragraph_before()
        run_pu = para_unico.add_run(
            'PARÁGRAFO ÚNICO - No caso de atraso de pagamento, o desconto será suprimido do saldo '
            'remanescente das mensalidades até que haja a regularização das parcelas em aberto, '
            'sem o prejuízo das demais cláusulas contratuais. Após a regularização das parcelas em aberto, '
            'o desconto poderá retornar incidindo única e exclusivamente sobre o saldo remanescente '
            'da anuidade a partir da parcela do mês subsequente a aquela data.'
        )
        
        if ref_para.style:
            para_unico.style = ref_para.style
        
        print("✅ Cláusula de desconto inserida com tags íntegras")
        
        # Salva o documento
        print(f"💾 Salvando: {dest}")
        doc.save(dest)
        
        print("✅ Template de desconto criado com sucesso!")
        print("🔍 Tags inseridas: {{desconto}} e {{desconto_extenso}}")
        return True
        
    except Exception as e:
        print(f"❌ Erro: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == '__main__':
    print("=" * 70)
    print("FIX DISCOUNT TEMPLATE - Versão Limpa")
    print("=" * 70)
    
    if create_clean_discount_template():
        print("=" * 70)
        print("✅ SUCESSO! Template pronto para uso.")
        print("=" * 70)
    else:
        print("=" * 70)
        print("❌ FALHA na criação do template")
        print("=" * 70)
