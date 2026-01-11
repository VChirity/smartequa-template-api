from docx import Document
import shutil
import os

def add_tags_to_template(source_path, dest_path, template_type):
    """
    Adiciona tags aos templates originais preservando TODA a formatação, logo e estrutura
    """
    print(f"\n🔧 Processando: {source_path}")
    
    if not os.path.exists(source_path):
        print(f"❌ Arquivo não encontrado!")
        return False
    
    try:
        # Carrega o documento original
        doc = Document(source_path)
        print(f"✅ Documento original carregado")
        
        # Procura por textos específicos e substitui por tags
        # Vamos procurar padrões comuns e substituir
        
        replacements = {
            # Padrões que podem estar no documento original
            'NOME DO RESPONSÁVEL': '{{responsavel1}}',
            'Nome do Responsável': '{{responsavel1}}',
            'nome do responsável': '{{responsavel1}}',
            
            'CPF DO RESPONSÁVEL': '{{cpf_responsavel}}',
            'CPF do Responsável': '{{cpf_responsavel}}',
            'cpf do responsável': '{{cpf_responsavel}}',
            
            'NATURALIDADE DO RESPONSÁVEL': '{{naturalidade_resp1}}',
            'Naturalidade do Responsável': '{{naturalidade_resp1}}',
            'naturalidade do responsável': '{{naturalidade_resp1}}',
            
            'DATA DE NASCIMENTO DO RESPONSÁVEL': '{{nasc_resp1}}',
            'Data de Nascimento do Responsável': '{{nasc_resp1}}',
            'data de nascimento do responsável': '{{nasc_resp1}}',
            
            'NOME DO ALUNO': '{{nome_aluno}}',
            'Nome do Aluno': '{{nome_aluno}}',
            'nome do aluno': '{{nome_aluno}}',
            
            'CPF DO ALUNO': '{{cpf_aluno}}',
            'CPF do Aluno': '{{cpf_aluno}}',
            'cpf do aluno': '{{cpf_aluno}}',
            
            'NATURALIDADE DO ALUNO': '{{naturalidade_aluno}}',
            'Naturalidade do Aluno': '{{naturalidade_aluno}}',
            'naturalidade do aluno': '{{naturalidade_aluno}}',
            
            'DATA DE NASCIMENTO DO ALUNO': '{{nasc_aluno}}',
            'Data de Nascimento do Aluno': '{{nasc_aluno}}',
            'data de nascimento do aluno': '{{nasc_aluno}}',
            
            'DATA POR EXTENSO': '{{data_extenso}}',
            'Data por Extenso': '{{data_extenso}}',
            'data por extenso': '{{data_extenso}}',
        }
        
        changes_made = 0
        
        # Processa parágrafos preservando formatação
        for para in doc.paragraphs:
            for run in para.runs:
                original_text = run.text
                new_text = original_text
                
                for old_text, tag in replacements.items():
                    if old_text in new_text:
                        new_text = new_text.replace(old_text, tag)
                        changes_made += 1
                
                if new_text != original_text:
                    run.text = new_text
        
        # Processa tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            original_text = run.text
                            new_text = original_text
                            
                            for old_text, tag in replacements.items():
                                if old_text in new_text:
                                    new_text = new_text.replace(old_text, tag)
                                    changes_made += 1
                            
                            if new_text != original_text:
                                run.text = new_text
        
        # Processa cabeçalhos e rodapés
        for section in doc.sections:
            # Header
            for para in section.header.paragraphs:
                for run in para.runs:
                    original_text = run.text
                    new_text = original_text
                    
                    for old_text, tag in replacements.items():
                        if old_text in new_text:
                            new_text = new_text.replace(old_text, tag)
                            changes_made += 1
                    
                    if new_text != original_text:
                        run.text = new_text
            
            # Footer
            for para in section.footer.paragraphs:
                for run in para.runs:
                    original_text = run.text
                    new_text = original_text
                    
                    for old_text, tag in replacements.items():
                        if old_text in new_text:
                            new_text = new_text.replace(old_text, tag)
                            changes_made += 1
                    
                    if new_text != original_text:
                        run.text = new_text
        
        # Salva o documento
        doc.save(dest_path)
        
        print(f"✅ Documento salvo: {dest_path}")
        print(f"📊 {changes_made} substituições realizadas")
        print(f"🎨 Formatação, logo e estrutura PRESERVADAS")
        
        return True
        
    except Exception as e:
        print(f"❌ Erro: {e}")
        import traceback
        traceback.print_exc()
        return False

def main():
    print("=" * 70)
    print("ADD TAGS TO ORIGINAL TEMPLATES")
    print("Preservando TODA formatação, logo e estrutura")
    print("=" * 70)
    
    templates = [
        {
            'source': 'assets/IMAGEM-E-VOZ-ALUNO-PUBLICIDADE_1.docx',
            'dest': 'templates/IMAGEM-E-VOZ-ALUNO-PUBLICIDADE_1.docx',
            'type': 'publicidade'
        },
        {
            'source': 'assets/IMAGEM-E-VOZ-ALUNO-INSTITUCIONAL_1.docx',
            'dest': 'templates/IMAGEM-E-VOZ-ALUNO-INSTITUCIONAL_1.docx',
            'type': 'institucional'
        }
    ]
    
    success_count = 0
    for template in templates:
        if add_tags_to_template(template['source'], template['dest'], template['type']):
            success_count += 1
    
    print("\n" + "=" * 70)
    print(f"✅ Concluído! {success_count}/{len(templates)} templates processados")
    print("=" * 70)
    
    if success_count == len(templates):
        print("\n🎉 Templates prontos com tags E formatação original!")
        print("💡 Abra os arquivos no Word para verificar.")
        print("📝 Se alguma tag não foi inserida, me avise que eu ajusto.")
    else:
        print("\n⚠️ Alguns templates falharam.")

if __name__ == '__main__':
    main()
