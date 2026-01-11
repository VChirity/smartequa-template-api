from docx import Document
import os

# Carregar o template original do usuário
template_original = 'templates_quadros/notas/quadro_notas_template.docx'

print(f"Carregando template original: {template_original}")
doc = Document(template_original)

# Procurar a tabela principal (geralmente a maior tabela do documento)
tabelas = doc.tables
print(f"Número de tabelas encontradas: {len(tabelas)}")

# Encontrar a tabela de notas (provavelmente a que tem mais colunas)
tabela_notas = None
for i, table in enumerate(tabelas):
    num_cols = len(table.columns)
    num_rows = len(table.rows)
    print(f"Tabela {i}: {num_rows} linhas x {num_cols} colunas")
    if num_cols >= 8:  # Tabela de notas tem 8 colunas
        tabela_notas = table
        print(f"  -> Esta parece ser a tabela de notas!")

if not tabela_notas:
    print("ERRO: Tabela de notas não encontrada!")
    exit(1)

print(f"\nTabela de notas encontrada com {len(tabela_notas.rows)} linhas")

# Procurar as linhas com tags do loop
linha_for = None
linha_dados = None
linha_endfor = None

for i, row in enumerate(tabela_notas.rows):
    texto_linha = ' '.join([cell.text for cell in row.cells])
    
    if 'for aluno in alunos' in texto_linha.lower():
        linha_for = i
        print(f"Linha {i}: Encontrado início do loop")
    elif 'endfor' in texto_linha.lower():
        linha_endfor = i
        print(f"Linha {i}: Encontrado fim do loop")
    elif '{{aluno.nome}}' in texto_linha or '{{loop.index}}' in texto_linha:
        linha_dados = i
        print(f"Linha {i}: Encontrada linha de dados")

if linha_for is None or linha_dados is None or linha_endfor is None:
    print("\nERRO: Não encontrei as linhas do loop!")
    print("Procurando por tags alternativas...")
    
    for i, row in enumerate(tabela_notas.rows):
        print(f"Linha {i}: {row.cells[0].text[:50]}...")

    exit(1)

print(f"\nModificando linhas do loop:")
print(f"  - Linha {linha_for}: início do loop")
print(f"  - Linha {linha_dados}: dados")
print(f"  - Linha {linha_endfor}: fim do loop")

# Modificar linha de início do loop
row_for = tabela_notas.rows[linha_for]
cell_for = row_for.cells[0]
# Limpar e adicionar tag correta
cell_for.text = ''
cell_for.paragraphs[0].text = '{%tr for aluno in alunos %}'
print("  ✓ Linha de início modificada")

# Modificar linha de fim do loop
row_endfor = tabela_notas.rows[linha_endfor]
cell_endfor = row_endfor.cells[0]
# Limpar e adicionar tag correta
cell_endfor.text = ''
cell_endfor.paragraphs[0].text = '{%tr endfor %}'
print("  ✓ Linha de fim modificada")

# A linha de dados não precisa ser modificada, apenas garantir que as tags estão corretas
print("  ✓ Linha de dados mantida")

# Salvar como novo arquivo
output_path = 'templates_quadros/notas/quadro_notas_template_CORRIGIDO.docx'
doc.save(output_path)

print(f"\n✅ Template corrigido salvo em: {output_path}")
print("\nAgora teste com: python test_template.py")
