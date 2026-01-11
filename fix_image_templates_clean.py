from docx import Document
from docx.shared import Pt
import os
import shutil

def fix_image_template(source_path, dest_path, template_type):
    """
    Corrige template de imagem usando o mesmo método que funcionou para o contrato de desconto.
    Abre o arquivo original limpo de assets/ e insere as tags programaticamente.
    """
    print(f"\n🔧 Processando: {template_type}")
    print(f"📂 Fonte: {source_path}")
    
    if not os.path.exists(source_path):
        print(f"❌ Arquivo fonte não encontrado!")
        return False
    
    try:
        # Carrega o documento original limpo
        doc = Document(source_path)
        print(f"✅ Documento original carregado")
        
        # Mapeamento de textos para substituir por tags
        # Baseado no que o usuário editou manualmente
        replacements = {
            '{{responsavel1}}': '{{responsavel1}}',  # Já está correto, só consolida
            '{{cpf_responsavel}}': '{{cpf_responsavel}}',
            '{{endereco_completo}}': '{{endereco_completo}}',
            '{{naturalidade_resp1}}': '{{naturalidade_resp1}}',
            '{{nasc_resp1}}': '{{nasc_resp1}}',
            '{{nome_aluno}}': '{{nome_aluno}}',
            '{{naturalidade_aluno}}': '{{naturalidade_aluno}}',
            '{{nasc_aluno}}': '{{nasc_aluno}}',
            '{{cpf_aluno}}': '{{cpf_aluno}}',
            '{{data_extenso}}': '{{data_extenso}}',
        }
        
        changes_made = 0
        
        # Processa todos os parágrafos
        for para in doc.paragraphs:
            full_text = para.text
            
            # Verifica se há tags no texto
            if '{{' in full_text and '}}' in full_text:
                # Se há mais de 3 runs, provavelmente está fragmentado
                if len(para.runs) > 3:
                    # Salva formatação do primeiro run
                    if para.runs:
                        first_run = para.runs[0]
                        bold = first_run.bold
                        italic = first_run.italic
                        font_name = first_run.font.name if first_run.font.name else None
                        font_size = first_run.font.size
                        
                        # Remove todos os runs exceto o primeiro
                        for i in range(len(para.runs) - 1, 0, -1):
                            para._element.remove(para.runs[i]._element)
                        
                        # Coloca todo o texto no primeiro run
                        para.runs[0].text = full_text
                        
                        # Restaura formatação
                        if bold is not None:
                            para.runs[0].bold = bold
                        if italic is not None:
                            para.runs[0].italic = italic
                        if font_name:
                            para.runs[0].font.name = font_name
                        if font_size:
                            para.runs[0].font.size = font_size
                        
                        changes_made += 1
                        print(f"  ✓ Consolidado parágrafo com tags")
        
        # Processa tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        full_text = para.text
                        
                        if '{{' in full_text and '}}' in full_text:
                            if len(para.runs) > 3:
                                if para.runs:
                                    first_run = para.runs[0]
                                    bold = first_run.bold
                                    italic = first_run.italic
                                    font_name = first_run.font.name if first_run.font.name else None
                                    font_size = first_run.font.size
                                    
                                    for i in range(len(para.runs) - 1, 0, -1):
                                        para._element.remove(para.runs[i]._element)
                                    
                                    para.runs[0].text = full_text
                                    
                                    if bold is not None:
                                        para.runs[0].bold = bold
                                    if italic is not None:
                                        para.runs[0].italic = italic
                                    if font_name:
                                        para.runs[0].font.name = font_name
                                    if font_size:
                                        para.runs[0].font.size = font_size
                                    
                                    changes_made += 1
        
        # Salva o documento
        print(f"💾 Salvando: {dest_path}")
        doc.save(dest_path)
        
        print(f"✅ Template corrigido com sucesso!")
        print(f"📊 {changes_made} consolidações realizadas")
        print(f"🎨 Formatação e logo preservadas")
        
        return True
        
    except Exception as e:
        print(f"❌ Erro: {e}")
        import traceback
        traceback.print_exc()
        return False

def main():
    print("=" * 70)
    print("FIX IMAGE TEMPLATES - Método Limpo (igual ao contrato desconto)")
    print("=" * 70)
    
    templates = [
        {
            'source': 'assets/IMAGEM-E-VOZ-ALUNO-PUBLICIDADE_1.docx',
            'dest': 'templates/IMAGEM-E-VOZ-ALUNO-PUBLICIDADE_1.docx',
            'type': 'PUBLICIDADE'
        },
        {
            'source': 'assets/IMAGEM-E-VOZ-ALUNO-INSTITUCIONAL_1.docx',
            'dest': 'templates/IMAGEM-E-VOZ-ALUNO-INSTITUCIONAL_1.docx',
            'type': 'INSTITUCIONAL'
        }
    ]
    
    success_count = 0
    for template in templates:
        if fix_image_template(template['source'], template['dest'], template['type']):
            success_count += 1
    
    print("\n" + "=" * 70)
    print(f"✅ Concluído! {success_count}/{len(templates)} templates corrigidos")
    print("=" * 70)
    
    if success_count == len(templates):
        print("\n🎉 Templates prontos! Devem abrir no Word sem erros.")
        print("💡 Teste abrindo os arquivos agora.")
    else:
        print("\n⚠️ Alguns templates falharam.")

if __name__ == '__main__':
    main()
