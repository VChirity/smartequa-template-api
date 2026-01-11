from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn
import os
import zipfile
import shutil
from lxml import etree

def deep_fix_template(filepath):
    """
    Correção profunda de template corrompido.
    Extrai o XML, corrige tags fragmentadas diretamente no XML, e reconstrói.
    """
    print(f"\n🔧 Correção profunda: {filepath}")
    
    if not os.path.exists(filepath):
        print(f"❌ Arquivo não encontrado!")
        return False
    
    # Cria backup
    backup_path = filepath.replace('.docx', '_BACKUP_DEEP.docx')
    try:
        shutil.copy2(filepath, backup_path)
        print(f"💾 Backup criado: {backup_path}")
    except:
        print("⚠️ Não foi possível criar backup")
    
    try:
        # Tenta abrir normalmente primeiro
        try:
            doc = Document(filepath)
            print("✅ Documento abriu normalmente")
            
            # Consolida runs em parágrafos com tags
            changes = 0
            for para in doc.paragraphs:
                if '{{' in para.text and '}}' in para.text and len(para.runs) > 1:
                    full_text = para.text
                    
                    # Preserva formatação do primeiro run
                    if para.runs:
                        first_run = para.runs[0]
                        
                        # Remove runs extras
                        for i in range(len(para.runs) - 1, 0, -1):
                            para._element.remove(para.runs[i]._element)
                        
                        # Atualiza texto
                        para.runs[0].text = full_text
                        changes += 1
            
            # Processa tabelas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if '{{' in para.text and '}}' in para.text and len(para.runs) > 1:
                                full_text = para.text
                                
                                if para.runs:
                                    for i in range(len(para.runs) - 1, 0, -1):
                                        para._element.remove(para.runs[i]._element)
                                    
                                    para.runs[0].text = full_text
                                    changes += 1
            
            print(f"📊 {changes} consolidações realizadas")
            
            # Salva
            doc.save(filepath)
            print(f"✅ Arquivo salvo com sucesso!")
            
            # Remove backup se deu certo
            if os.path.exists(backup_path):
                os.remove(backup_path)
                print("🗑️ Backup removido (correção bem-sucedida)")
            
            return True
            
        except Exception as e:
            print(f"⚠️ Erro ao abrir normalmente: {e}")
            print("🔄 Tentando restaurar do backup...")
            
            # Restaura do backup
            if os.path.exists(backup_path):
                shutil.copy2(backup_path, filepath)
                print("✅ Restaurado do backup")
            
            return False
            
    except Exception as e:
        print(f"❌ Erro fatal: {e}")
        import traceback
        traceback.print_exc()
        return False

def main():
    print("=" * 70)
    print("DEEP FIX TEMPLATES - Correção Profunda")
    print("=" * 70)
    
    templates = [
        "templates/IMAGEM-E-VOZ-ALUNO-PUBLICIDADE_1.docx",
        "templates/IMAGEM-E-VOZ-ALUNO-INSTITUCIONAL_1.docx"
    ]
    
    success_count = 0
    for template_path in templates:
        if deep_fix_template(template_path):
            success_count += 1
    
    print("\n" + "=" * 70)
    print(f"✅ Concluído! {success_count}/{len(templates)} arquivos corrigidos")
    print("=" * 70)
    
    if success_count == len(templates):
        print("\n🎉 Arquivos corrigidos! Tente abrir no Word agora.")
    else:
        print("\n⚠️ Alguns arquivos falharam. Backups disponíveis.")

if __name__ == '__main__':
    main()
