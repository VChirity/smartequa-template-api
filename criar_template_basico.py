from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_border(cell, **kwargs):
    """Adicionar bordas a uma célula"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_element = OxmlElement(f'w:{edge}')
        edge_element.set(qn('w:val'), 'single')
        edge_element.set(qn('w:sz'), '4')
        edge_element.set(qn('w:space'), '0')
        edge_element.set(qn('w:color'), '000000')
        tcBorders.append(edge_element)
    tcPr.append(tcBorders)

def set_cell_background(cell, color):
    """Definir cor de fundo da célula"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

# Criar documento
doc = Document()

# Configurar margens
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

# Título
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('QUADRO DE NOTAS')
run.bold = True
run.font.size = Pt(16)

# Subtítulo com tags
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('{{bimestre}} — {{ano}}')
run.font.size = Pt(12)

# Espaço
doc.add_paragraph()

# Turma
turma_para = doc.add_paragraph()
turma_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = turma_para.add_run('{{turma}}')
run.bold = True
run.font.size = Pt(14)
set_cell_bg = lambda p, color: None  # Placeholder

# Espaço
doc.add_paragraph()

# Tabela de notas
table = doc.add_table(rows=4, cols=8)
table.style = 'Table Grid'

# Linha 0: Cabeçalho "NOTAS"
row0 = table.rows[0]
cell = row0.cells[0]
for i in range(1, 8):
    cell.merge(row0.cells[i])
cell.text = 'NOTAS'
cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
run = cell.paragraphs[0].runs[0]
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(255, 255, 255)
set_cell_background(cell, '1F4E78')  # Azul escuro

# Linha 1: Cabeçalhos das colunas
headers = ['Nº', 'NOME DO ALUNO', '{{tipo_av_1}}\n{{pont_1}}', '{{tipo_av_2}}\n{{pont_2}}', 
           '{{tipo_av_3}}\n{{pont_3}}', '{{tipo_av_4}}\n{{pont_4}}', '{{tipo_av_5}}\n{{pont_5}}', 'MÉDIA']
row1 = table.rows[1]
for i, header in enumerate(headers):
    cell = row1.cells[i]
    cell.text = header
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cell.paragraphs[0].runs[0]
    run.bold = True
    run.font.size = Pt(10)
    if i < 2:
        set_cell_background(cell, '5B9BD5')  # Azul claro
    elif i == 7:
        set_cell_background(cell, '70AD47')  # Verde
    else:
        set_cell_background(cell, 'FFC000')  # Laranja

# Linha 2: Início do loop (mesclada)
row2 = table.rows[2]
cell = row2.cells[0]
for i in range(1, 8):
    cell.merge(row2.cells[i])
cell.text = '{% for aluno in alunos -%}'
cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
run = cell.paragraphs[0].runs[0]
run.font.size = Pt(9)
set_cell_background(cell, 'D9E1F2')  # Azul muito claro

# Linha 3: Dados dos alunos
row3 = table.rows[3]
dados = ['{{loop.index}}', '{{aluno.nome}}', '{{aluno.av1}}', '{{aluno.av2}}', 
         '{{aluno.av3}}', '{{aluno.av4}}', '{{aluno.av5}}', '{{aluno.media}}']
for i, dado in enumerate(dados):
    cell = row3.cells[i]
    cell.text = dado
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cell.paragraphs[0].runs[0]
    run.font.size = Pt(10)

# Adicionar linha para fim do loop
row4 = table.add_row()
cell = row4.cells[0]
for i in range(1, 8):
    cell.merge(row4.cells[i])
cell.text = '{% endfor -%}'
cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
run = cell.paragraphs[0].runs[0]
run.font.size = Pt(9)
set_cell_background(cell, 'D9E1F2')  # Azul muito claro

# Espaço
doc.add_paragraph()
doc.add_paragraph()

# Cálculo da média
doc.add_paragraph('CÁLCULO DA MÉDIA')
calculo_para = doc.add_paragraph()
calculo_para.add_run('☐ Nota única    ☐ Média Aritmética    ☐ Adição das notas    ☐ Outro')
doc.add_paragraph('Como calcular a média: {{explicacao_calculo}}')

# Espaço
doc.add_paragraph()

# Rodapé
footer_para = doc.add_paragraph()
footer_para.add_run('Professor: {{professor}}')
footer_para = doc.add_paragraph()
footer_para.add_run('Disciplina: {{disciplina}}')

# Salvar
output_path = 'templates_quadros/notas/quadro_notas_template_NOVO.docx'
doc.save(output_path)
print(f'✅ Template criado: {output_path}')
print('Agora você pode:')
print('1. Abrir o arquivo e adicionar sua logo')
print('2. Ajustar cores conforme preferir')
print('3. Renomear para quadro_notas_template.docx (substituir o antigo)')
print('4. Testar novamente')
