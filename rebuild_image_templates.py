from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_image_template(template_type):
    """
    Cria template de termo de imagem do zero com estrutura limpa
    """
    
    if template_type == "publicidade":
        dest = "templates/IMAGEM-E-VOZ-ALUNO-PUBLICIDADE_1.docx"
        titulo = "TERMO DE AUTORIZAÇÃO DE USO DE IMAGEM E VOZ - PUBLICIDADE"
    else:
        dest = "templates/IMAGEM-E-VOZ-ALUNO-INSTITUCIONAL_1.docx"
        titulo = "TERMO DE AUTORIZAÇÃO DE USO DE IMAGEM E VOZ - INSTITUCIONAL"
    
    print(f"\n🔧 Criando: {dest}")
    
    try:
        # Cria novo documento
        doc = Document()
        
        # Configurações de página
        section = doc.sections[0]
        section.page_height = Inches(11.69)  # A4
        section.page_width = Inches(8.27)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        
        # Título
        titulo_para = doc.add_paragraph()
        titulo_run = titulo_para.add_run(titulo)
        titulo_run.bold = True
        titulo_run.font.size = Pt(14)
        titulo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Espaço
        
        # Texto do termo
        if template_type == "publicidade":
            texto = (
                "Eu, "
            )
            doc.add_paragraph(texto)
            
            # Parágrafo com tags (cada tag em seu próprio run)
            p1 = doc.add_paragraph()
            p1.add_run("Eu, ")
            p1.add_run("{{responsavel1}}").bold = True
            p1.add_run(", portador(a) do CPF ")
            p1.add_run("{{cpf_responsavel}}").bold = True
            p1.add_run(", natural de ")
            p1.add_run("{{naturalidade_resp1}}").bold = True
            p1.add_run(", nascido(a) em ")
            p1.add_run("{{nasc_resp1}}").bold = True
            p1.add_run(", responsável legal pelo(a) menor ")
            p1.add_run("{{nome_aluno}}").bold = True
            p1.add_run(", CPF ")
            p1.add_run("{{cpf_aluno}}").bold = True
            p1.add_run(", natural de ")
            p1.add_run("{{naturalidade_aluno}}").bold = True
            p1.add_run(", nascido(a) em ")
            p1.add_run("{{nasc_aluno}}").bold = True
            p1.add_run(", AUTORIZO o uso da imagem e voz do(a) menor para fins de PUBLICIDADE e MARKETING da instituição Smart Equação, incluindo redes sociais, materiais impressos e campanhas publicitárias.")
            
            doc.add_paragraph()
            
            p2 = doc.add_paragraph()
            p2.add_run("Esta autorização é válida por tempo indeterminado e pode ser revogada a qualquer momento mediante comunicação por escrito.")
            
        else:  # institucional
            p1 = doc.add_paragraph()
            p1.add_run("Eu, ")
            p1.add_run("{{responsavel1}}").bold = True
            p1.add_run(", portador(a) do CPF ")
            p1.add_run("{{cpf_responsavel}}").bold = True
            p1.add_run(", natural de ")
            p1.add_run("{{naturalidade_resp1}}").bold = True
            p1.add_run(", nascido(a) em ")
            p1.add_run("{{nasc_resp1}}").bold = True
            p1.add_run(", responsável legal pelo(a) menor ")
            p1.add_run("{{nome_aluno}}").bold = True
            p1.add_run(", CPF ")
            p1.add_run("{{cpf_aluno}}").bold = True
            p1.add_run(", natural de ")
            p1.add_run("{{naturalidade_aluno}}").bold = True
            p1.add_run(", nascido(a) em ")
            p1.add_run("{{nasc_aluno}}").bold = True
            p1.add_run(", AUTORIZO o uso da imagem e voz do(a) menor para fins INSTITUCIONAIS da Smart Equação, incluindo eventos internos, materiais pedagógicos e registros acadêmicos.")
            
            doc.add_paragraph()
            
            p2 = doc.add_paragraph()
            p2.add_run("Esta autorização é válida durante o período de matrícula do(a) aluno(a) e pode ser revogada a qualquer momento mediante comunicação por escrito.")
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Assinatura
        p_data = doc.add_paragraph()
        p_data.add_run("Rio de Janeiro, ")
        p_data.add_run("{{data_extenso}}").bold = True
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        p_ass = doc.add_paragraph("_" * 50)
        p_ass.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p_nome = doc.add_paragraph()
        p_nome.add_run("{{responsavel1}}").bold = True
        p_nome.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p_cpf = doc.add_paragraph()
        p_cpf.add_run("CPF: ")
        p_cpf.add_run("{{cpf_responsavel}}").bold = True
        p_cpf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Salva
        doc.save(dest)
        print(f"✅ Template criado com sucesso!")
        print(f"🔍 Tags inseridas: responsavel1, cpf_responsavel, naturalidade_resp1, nasc_resp1")
        print(f"   nome_aluno, cpf_aluno, naturalidade_aluno, nasc_aluno, data_extenso")
        return True
        
    except Exception as e:
        print(f"❌ Erro: {e}")
        import traceback
        traceback.print_exc()
        return False

def main():
    print("=" * 70)
    print("REBUILD IMAGE TEMPLATES - Reconstrução Completa")
    print("=" * 70)
    
    success_count = 0
    
    if create_image_template("publicidade"):
        success_count += 1
    
    if create_image_template("institucional"):
        success_count += 1
    
    print("\n" + "=" * 70)
    print(f"✅ Concluído! {success_count}/2 templates criados")
    print("=" * 70)
    
    if success_count == 2:
        print("\n🎉 Templates reconstruídos do zero!")
        print("💡 Abra os arquivos no Word para verificar.")
        print("📝 Você pode editar o conteúdo diretamente no Word se necessário.")
    else:
        print("\n⚠️ Alguns templates falharam na criação.")

if __name__ == '__main__':
    main()
