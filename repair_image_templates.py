from docx import Document
import os
import shutil

def repair_template(filepath):
    """
    Tenta reparar template corrompido verificando e corrigindo apenas tags fragmentadas
    de forma mais conservadora, preservando a estrutura do documento
    """
    if not os.path.exists(filepath):
        print(f"❌ Arquivo não encontrado: {filepath}")
        return False
    
    print(f"\n🔧 Reparando: {filepath}")
    
    # Cria backup antes de modificar
    backup_path = filepath.replace('.docx', '_BACKUP.docx')
    try:
        shutil.copy2(filepath, backup_path)
        print(f"💾 Backup criado: {backup_path}")
    except:
        print("⚠️ Não foi possível criar backup")
    
    try:
        doc = Document(filepath)
        changes_made = 0
        
        # Processa parágrafos de forma mais conservadora
        for para in doc.paragraphs:
            if '{{' in para.text and '}}' in para.text:
                # Verifica se a tag está fragmentada
                full_text = para.text
                
                # Conta quantos runs existem
                num_runs = len(para.runs)
                
                # Se há muitos runs em um parágrafo curto, pode estar fragmentado
                if num_runs > 5 and len(full_text) < 200:
                    # Preserva formatação do primeiro run
                    if para.runs:
                        first_run = para.runs[0]
                        bold = first_run.bold
                        italic = first_run.italic
                        font_name = first_run.font.name if first_run.font.name else None
                        font_size = first_run.font.size
                        
                        # Limpa runs
                        for run in para.runs[1:]:
                            run.text = ''
                        
                        # Coloca texto no primeiro run
                        first_run.text = full_text
                        
                        # Restaura formatação
                        if bold is not None:
                            first_run.bold = bold
                        if italic is not None:
                            first_run.italic = italic
                        if font_name:
                            first_run.font.name = font_name
                        if font_size:
                            first_run.font.size = font_size
                        
                        changes_made += 1
        
        # Processa tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if '{{' in para.text and '}}' in para.text:
                            full_text = para.text
                            num_runs = len(para.runs)
                            
                            if num_runs > 5 and len(full_text) < 200:
                                if para.runs:
                                    first_run = para.runs[0]
                                    bold = first_run.bold
                                    italic = first_run.italic
                                    font_name = first_run.font.name if first_run.font.name else None
                                    font_size = first_run.font.size
                                    
                                    for run in para.runs[1:]:
                                        run.text = ''
                                    
                                    first_run.text = full_text
                                    
                                    if bold is not None:
                                        first_run.bold = bold
                                    if italic is not None:
                                        first_run.italic = italic
                                    if font_name:
                                        first_run.font.name = font_name
                                    if font_size:
                                        first_run.font.size = font_size
                                    
                                    changes_made += 1
        
        if changes_made > 0:
            doc.save(filepath)
            print(f"✅ Reparado! {changes_made} parágrafos corrigidos")
            print(f"📁 Backup mantido em: {backup_path}")
            return True
        else:
            print(f"ℹ️ Nenhuma correção necessária")
            # Remove backup se não houve mudanças
            if os.path.exists(backup_path):
                os.remove(backup_path)
            return True
            
    except Exception as e:
        print(f"❌ Erro ao reparar: {e}")
        print(f"💡 Restaurando do backup...")
        
        # Tenta restaurar do backup
        if os.path.exists(backup_path):
            try:
                shutil.copy2(backup_path, filepath)
                print(f"✅ Arquivo restaurado do backup")
            except:
                print(f"❌ Falha ao restaurar backup")
        
        import traceback
        traceback.print_exc()
        return False

def main():
    print("=" * 70)
    print("REPAIR IMAGE TEMPLATES - Correção Conservadora")
    print("=" * 70)
    
    templates = [
        "templates/IMAGEM-E-VOZ-ALUNO-PUBLICIDADE_1.docx",
        "templates/IMAGEM-E-VOZ-ALUNO-INSTITUCIONAL_1.docx"
    ]
    
    success_count = 0
    for template_path in templates:
        if repair_template(template_path):
            success_count += 1
    
    print("\n" + "=" * 70)
    print(f"✅ Concluído! {success_count}/{len(templates)} arquivos processados")
    print("=" * 70)
    
    if success_count == len(templates):
        print("\n🎉 Tente abrir os arquivos agora no Word!")
        print("💡 Se ainda houver problemas, os backups estão disponíveis.")
    else:
        print("\n⚠️ Alguns arquivos podem precisar de atenção manual.")
        print("💡 Verifique os backups criados.")

if __name__ == '__main__':
    main()
