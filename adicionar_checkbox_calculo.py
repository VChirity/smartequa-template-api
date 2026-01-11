from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Carregar o template atual
template_path = 'templates_quadros/notas/quadro_notas_template.docx'

print(f"Carregando template: {template_path}")
doc = Document(template_path)

# Procurar o parágrafo com os checkboxes
encontrado = False
for i, para in enumerate(doc.paragraphs):
    texto = para.text
    
    # Procurar pela linha dos checkboxes
    if 'Nota única' in texto or 'Média Aritmética' in texto or 'Adição das notas' in texto:
        print(f"\nParágrafo {i} encontrado:")
        print(f"Texto atual: {texto}")
        
        # Limpar o parágrafo
        para.clear()
        
        # Adicionar texto com lógica condicional Jinja2
        novo_texto = (
            "{% if tipo_calculo == 'Nota única' %}☑{% else %}☐{% endif %} Nota única    "
            "{% if tipo_calculo == 'Média Aritmética' %}☑{% else %}☐{% endif %} Média Aritmética    "
            "{% if tipo_calculo == 'Adição das notas' %}☑{% else %}☐{% endif %} Adição das notas    "
            "{% if tipo_calculo == 'Outro' %}☑{% else %}☐{% endif %} Outro"
        )
        
        para.add_run(novo_texto)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        print(f"Novo texto: {novo_texto}")
        encontrado = True
        break

if not encontrado:
    print("\nAVISO: Não encontrei o parágrafo dos checkboxes!")
    print("Procurando em todas as tabelas...")
    
    # Procurar em tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    texto = para.text
                    if 'Nota única' in texto or 'Média Aritmética' in texto:
                        print(f"\nEncontrado em tabela:")
                        print(f"Texto atual: {texto}")
                        
                        # Limpar e adicionar novo texto
                        para.clear()
                        novo_texto = (
                            "{% if tipo_calculo == 'Nota única' %}☑{% else %}☐{% endif %} Nota única    "
                            "{% if tipo_calculo == 'Média Aritmética' %}☑{% else %}☐{% endif %} Média Aritmética    "
                            "{% if tipo_calculo == 'Adição das notas' %}☑{% else %}☐{% endif %} Adição das notas    "
                            "{% if tipo_calculo == 'Outro' %}☑{% else %}☐{% endif %} Outro"
                        )
                        para.add_run(novo_texto)
                        print(f"Novo texto: {novo_texto}")
                        encontrado = True
                        break
                if encontrado:
                    break
            if encontrado:
                break
        if encontrado:
            break

if encontrado:
    # Salvar
    output_path = 'templates_quadros/notas/quadro_notas_template.docx'
    doc.save(output_path)
    print(f"\n✅ Template atualizado com checkboxes condicionais!")
    print("Teste novamente no SmartEqua!")
else:
    print("\n❌ ERRO: Não consegui encontrar a linha dos checkboxes!")
    print("Listando todos os parágrafos para debug:")
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f"  {i}: {para.text[:100]}")
