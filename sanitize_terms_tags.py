from docx import Document
import os

def sanitize_tags(filepath):
    """
    Sanitiza tags nos arquivos de termo de imagem.
    Remove acentos, corrige nomes de variáveis e força reescrita do XML
    para unificar runs fragmentados.
    """
    print(f"\n🔧 Sanitizando: {filepath}")
    
    if not os.path.exists(filepath):
        print(f"❌ Arquivo não encontrado!")
        return False
    
    try:
        doc = Document(filepath)
        print(f"✅ Documento carregado")
        
        changes_made = 0
        
        # Mapeamento de substituições para corrigir acentos e nomes
        replacements = {
            '{{endereço': '{{endereco',
            '{{responsável': '{{responsavel',
            '{{naturalidade_responsável': '{{naturalidade_resp1',
            '{{nasc_responsável': '{{nasc_resp1',
            # Garante integridade das tags corretas
            '{{endereco_completo}}': '{{endereco_completo}}',
            '{{responsavel1}}': '{{responsavel1}}',
            '{{cpf_responsavel}}': '{{cpf_responsavel}}',
            '{{naturalidade_resp1}}': '{{naturalidade_resp1}}',
            '{{nasc_resp1}}': '{{nasc_resp1}}',
            '{{nome_aluno}}': '{{nome_aluno}}',
            '{{naturalidade_aluno}}': '{{naturalidade_aluno}}',
            '{{nasc_aluno}}': '{{nasc_aluno}}',
            '{{cpf_aluno}}': '{{cpf_aluno}}',
            '{{data_extenso}}': '{{data_extenso}}',
        }
        
        # Processa todos os parágrafos
        for para in doc.paragraphs:
            text = para.text
            original_text = text
            
            # Aplica substituições
            for old, new in replacements.items():
                if old in text:
                    text = text.replace(old, new)
            
            # Se houve alteração OU se tem tags, reescreve o parágrafo
            # Isso força o python-docx a unificar runs fragmentados
            if text != original_text or '{{' in text:
                para.text = text
                changes_made += 1
                if text != original_text:
                    print(f"  ✓ Corrigido: {original_text[:50]}...")
                else:
                    print(f"  ✓ Unificado: {text[:50]}...")
        
        # Processa tabelas (se houver)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = para.text
                        original_text = text
                        
                        for old, new in replacements.items():
                            if old in text:
                                text = text.replace(old, new)
                        
                        if text != original_text or '{{' in text:
                            para.text = text
                            changes_made += 1
        
        # Processa cabeçalhos e rodapés
        for section in doc.sections:
            for para in section.header.paragraphs:
                text = para.text
                original_text = text
                
                for old, new in replacements.items():
                    if old in text:
                        text = text.replace(old, new)
                
                if text != original_text or '{{' in text:
                    para.text = text
                    changes_made += 1
            
            for para in section.footer.paragraphs:
                text = para.text
                original_text = text
                
                for old, new in replacements.items():
                    if old in text:
                        text = text.replace(old, new)
                
                if text != original_text or '{{' in text:
                    para.text = text
                    changes_made += 1
        
        # Salva o documento
        doc.save(filepath)
        
        print(f"✅ Documento salvo com {changes_made} parágrafos processados")
        print(f"🎨 Tags sanitizadas e runs unificados")
        
        return True
        
    except Exception as e:
        print(f"❌ Erro: {e}")
        import traceback
        traceback.print_exc()
        return False

def main():
    print("=" * 70)
    print("SANITIZE TERMS TAGS - Limpeza e Unificação de Tags")
    print("Remove acentos, corrige nomes e unifica runs fragmentados")
    print("=" * 70)
    
    templates = [
        "templates/IMAGEM-E-VOZ-ALUNO-PUBLICIDADE_1.docx",
        "templates/IMAGEM-E-VOZ-ALUNO-INSTITUCIONAL_1.docx"
    ]
    
    success_count = 0
    for template_path in templates:
        if sanitize_tags(template_path):
            success_count += 1
    
    print("\n" + "=" * 70)
    print(f"✅ Concluído! {success_count}/{len(templates)} arquivos sanitizados")
    print("=" * 70)
    
    if success_count == len(templates):
        print("\n🎉 Tags sanitizadas e unificadas!")
        print("💡 Método: Reescrita forçada do XML para unificar runs")
        print("🧪 Teste no app agora - deve funcionar sem corrupção!")
    else:
        print("\n⚠️ Alguns arquivos falharam.")

if __name__ == '__main__':
    main()
