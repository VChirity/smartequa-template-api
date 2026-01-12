from docx import Document
from docx.shared import RGBColor
import os

def aplicar_negrito_em_tags(doc_path):
    """
    Aplica negrito em todas as tags {{ }} no documento Word
    """
    print(f"\n🔧 Processando: {os.path.basename(doc_path)}")
    
    doc = Document(doc_path)
    tags_encontradas = 0
    
    # Processar parágrafos
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if '{{' in run.text or '}}' in run.text:
                run.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)  # Preto
                tags_encontradas += 1
                print(f"  ✅ Tag encontrada e formatada: {run.text[:50]}...")
    
    # Processar tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if '{{' in run.text or '}}' in run.text:
                            run.bold = True
                            run.font.color.rgb = RGBColor(0, 0, 0)
                            tags_encontradas += 1
                            print(f"  ✅ Tag em tabela formatada: {run.text[:50]}...")
    
    # Salvar documento
    doc.save(doc_path)
    print(f"  💾 Documento salvo com {tags_encontradas} tags formatadas em negrito")
    return tags_encontradas

if __name__ == '__main__':
    # Templates dos contratos
    templates = [
        'templates_contratos/template_contrato2025_2.docx',
        'templates_contratos/template_contratoDESCONTO2025_2.docx',
    ]
    
    print("=" * 60)
    print("🔨 APLICANDO NEGRITO NAS TAGS DOS CONTRATOS")
    print("=" * 60)
    
    total_tags = 0
    for template in templates:
        if os.path.exists(template):
            tags = aplicar_negrito_em_tags(template)
            total_tags += tags
        else:
            print(f"❌ Template não encontrado: {template}")
    
    print("\n" + "=" * 60)
    print(f"✅ CONCLUÍDO! Total de {total_tags} tags formatadas em negrito")
    print("=" * 60)
