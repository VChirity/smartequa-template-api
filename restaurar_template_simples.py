from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_background(cell, color):
    """Definir cor de fundo da célula"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

# Carregar template
template_path = 'templates_quadros/notas/quadro_notas_template.docx'
print(f"Carregando template: {template_path}")
doc = Document(template_path)

# Remover as últimas tabelas e parágrafos adicionados (seção de cálculo e rodapé)
print(f"Total de tabelas antes: {len(doc.tables)}")
print(f"Total de parágrafos antes: {len(doc.paragraphs)}")

# Manter apenas a primeira tabela (tabela de notas)
# Remover tabelas extras
while len(doc.tables) > 1:
    table = doc.tables[-1]
    table._element.getparent().remove(table._element)
    print(f"  Tabela removida")

# Remover parágrafos extras (manter apenas os 3 primeiros: título, subtítulo, turma)
# E os parágrafos vazios necessários
paragrafos_para_manter = []
for i, para in enumerate(doc.paragraphs):
    texto = para.text.strip()
    # Manter parágrafos importantes
    if 'QUADRO' in texto or '{{bimestre}}' in texto or '{{turma}}' in texto or texto == '':
        paragrafos_para_manter.append(para)
    else:
        # Remover parágrafos extras
        p_element = para._element
        p_element.getparent().remove(p_element)

print(f"Total de tabelas depois: {len(doc.tables)}")
print(f"Total de parágrafos depois: {len(doc.paragraphs)}")

# Adicionar espaços
doc.add_paragraph()
doc.add_paragraph()

# Adicionar tabela simples de cálculo
print("\nAdicionando tabela de cálculo simples...")
table_calculo = doc.add_table(rows=2, cols=1)

# Linha 0: Título
row0 = table_calculo.rows[0]
cell0 = row0.cells[0]
cell0.text = 'CÁLCULO DA MÉDIA'
cell0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
run = cell0.paragraphs[0].runs[0]
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(255, 255, 255)
set_cell_background(cell0, '1F4E78')

# Linha 1: Tipo de cálculo e explicação
row1 = table_calculo.rows[1]
cell1 = row1.cells[0]
set_cell_background(cell1, 'D9E1F2')
para1 = cell1.paragraphs[0]
para1.text = 'Tipo de cálculo: {{tipo_calculo}}\nComo calcular a média: {{explicacao_calculo}}'
para1.alignment = WD_ALIGN_PARAGRAPH.LEFT

print("✓ Tabela de cálculo adicionada")

# Adicionar espaço
doc.add_paragraph()

# Adicionar rodapé simples
print("Adicionando rodapé...")
para_prof = doc.add_paragraph()
run_prof = para_prof.add_run('Professor: {{professor}}')
run_prof.font.bold = True

para_disc = doc.add_paragraph()
run_disc = para_disc.add_run('Disciplina: {{disciplina}}')
run_disc.font.bold = True

# Aplicar fundo amarelo
for para in [para_prof, para_disc]:
    pPr = para._element.get_or_add_pPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'FFC000')
    pPr.append(shading)

print("✓ Rodapé adicionado")

# Salvar
output_path = 'templates_quadros/notas/quadro_notas_template.docx'
doc.save(output_path)

print(f"\n✅ Template restaurado e simplificado!")
print("Teste novamente no SmartEqua!")
