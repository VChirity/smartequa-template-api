from docx import Document

# Carregar template
template_path = 'templates_quadros/notas/quadro_notas_template.docx'
print(f"Carregando template: {template_path}")
doc = Document(template_path)

# Procurar a tabela de cálculo (segunda tabela)
if len(doc.tables) < 2:
    print("ERRO: Tabela de cálculo não encontrada!")
    exit(1)

tabela_calculo = doc.tables[1]
print(f"Tabela de cálculo encontrada com {len(tabela_calculo.rows)} linhas")

# Encontrar a célula com "Como calcular a média"
encontrado = False
for r_idx, row in enumerate(tabela_calculo.rows):
    for c_idx, cell in enumerate(row.cells):
        texto = cell.text
        if 'Como calcular a média' in texto or 'tipo_calculo' in texto:
            print(f"\nCélula encontrada na linha {r_idx}, célula {c_idx}")
            print(f"Texto atual: {texto[:100]}")
            
            # Limpar e adicionar texto com condicional
            cell.text = ''
            para = cell.paragraphs[0]
            
            # Novo texto com lógica condicional
            novo_texto = (
                "Tipo de cálculo: {{tipo_calculo}}\n"
                "{% if tipo_calculo == 'Outro' %}"
                "Como calcular a média: {{explicacao_calculo}}"
                "{% endif %}"
            )
            
            para.text = novo_texto
            print(f"Novo texto: {novo_texto}")
            encontrado = True
            break
    if encontrado:
        break

if encontrado:
    # Salvar
    output_path = 'templates_quadros/notas/quadro_notas_template.docx'
    doc.save(output_path)
    print(f"\n✅ Template atualizado!")
    print("Agora 'Como calcular a média' só aparece se tipo_calculo == 'Outro'")
else:
    print("\n❌ Não encontrei a célula com 'Como calcular a média'")
