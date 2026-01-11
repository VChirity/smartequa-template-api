from docx import Document
import os
import shutil

def extract_text_from_corrupted(corrupted_path):
    """
    Extrai o texto do arquivo corrompido (que o python-docx consegue abrir)
    """
    try:
        doc = Document(corrupted_path)
        paragraphs_text = []
        
        for para in doc.paragraphs:
            paragraphs_text.append(para.text)
        
        return paragraphs_text
    except Exception as e:
        print(f"Erro ao extrair texto: {e}")
        return None

def rebuild_from_clean_base(clean_base_path, corrupted_path, dest_path):
    """
    Reconstrói o template usando o arquivo limpo de assets/ como base
    e copiando o conteúdo do arquivo corrompido
    """
    print(f"\n🔧 Reconstruindo: {os.path.basename(dest_path)}")
    
    # Extrai texto do arquivo corrompido
    print("📖 Extraindo conteúdo do arquivo corrompido...")
    corrupted_text = extract_text_from_corrupted(corrupted_path)
    
    if not corrupted_text:
        print("❌ Falha ao extrair texto")
        return False
    
    print(f"✅ {len(corrupted_text)} parágrafos extraídos")
    
    # Carrega o arquivo limpo de assets/
    print("📂 Carregando arquivo limpo de assets/...")
    clean_doc = Document(clean_base_path)
    
    # Substitui o conteúdo dos parágrafos
    print("✏️ Substituindo conteúdo...")
    for i, para in enumerate(clean_doc.paragraphs):
        if i < len(corrupted_text):
            # Preserva formatação, apenas substitui o texto
            if para.runs:
                # Limpa runs existentes
                for run in para.runs[1:]:
                    para._element.remove(run._element)
                
                # Coloca o texto do arquivo corrompido no primeiro run
                para.runs[0].text = corrupted_text[i]
    
    # Salva o novo arquivo
    print(f"💾 Salvando: {dest_path}")
    clean_doc.save(dest_path)
    
    print(f"✅ Arquivo reconstruído com sucesso!")
    return True

def main():
    print("=" * 70)
    print("REBUILD FROM CORRUPTED - Reconstrução Total")
    print("Usa arquivos limpos de assets/ + conteúdo dos corrompidos")
    print("=" * 70)
    
    templates = [
        {
            'clean': 'assets/IMAGEM-E-VOZ-ALUNO-PUBLICIDADE_1.docx',
            'corrupted': 'templates backup/IMAGEM-E-VOZ-ALUNO-PUBLICIDADE_1.docx',
            'dest': 'templates/IMAGEM-E-VOZ-ALUNO-PUBLICIDADE_1.docx'
        },
        {
            'clean': 'assets/IMAGEM-E-VOZ-ALUNO-INSTITUCIONAL_1.docx',
            'corrupted': 'templates backup/IMAGEM-E-VOZ-ALUNO-INSTITUCIONAL_1.docx',
            'dest': 'templates/IMAGEM-E-VOZ-ALUNO-INSTITUCIONAL_1.docx'
        }
    ]
    
    success_count = 0
    for template in templates:
        if rebuild_from_clean_base(template['clean'], template['corrupted'], template['dest']):
            success_count += 1
    
    print("\n" + "=" * 70)
    print(f"✅ Concluído! {success_count}/{len(templates)} arquivos reconstruídos")
    print("=" * 70)
    
    if success_count == len(templates):
        print("\n🎉 Arquivos reconstruídos com estrutura limpa!")
        print("💡 Estrutura XML de assets/ (limpa) + conteúdo editado por você")
        print("🔍 Tente abrir no Word agora.")
    else:
        print("\n⚠️ Alguns arquivos falharam.")

if __name__ == '__main__':
    main()
